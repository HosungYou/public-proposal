#!/usr/bin/env python3
"""Audit an A4 Korean public-proposal DOCX against a JSON project profile."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def mm(emu: int | None) -> float | None:
    return None if emu is None else round(emu / 36000, 2)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--profile", type=Path, required=True)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()
    profile = json.loads(args.profile.read_text(encoding="utf-8"))
    doc = Document(args.docx)
    findings: list[dict] = []

    page = profile["page"]
    for idx, section in enumerate(doc.sections, 1):
        actual = {
            "width_mm": mm(section.page_width),
            "height_mm": mm(section.page_height),
            "top_mm": mm(section.top_margin),
            "bottom_mm": mm(section.bottom_margin),
            "left_mm": mm(section.left_margin),
            "right_mm": mm(section.right_margin),
        }
        for key, expected in page.items():
            if key in actual and abs(actual[key] - expected) > 0.6:
                findings.append({"rule": "page_geometry", "section": idx, "key": key, "expected": expected, "actual": actual[key], "severity": "blocker"})

    required_styles = profile.get("required_styles", [])
    style_names = {style.name for style in doc.styles}
    for style in required_styles:
        if style not in style_names:
            findings.append({"rule": "missing_style", "style": style, "severity": "blocker"})

    allowed_fonts = set(profile.get("allowed_fonts", []))
    for style_name in required_styles:
        if style_name not in style_names:
            continue
        style = doc.styles[style_name]
        fonts = style.element.get_or_add_rPr().find(qn("w:rFonts"))
        values = set()
        if fonts is not None:
            for attr in ("ascii", "hAnsi", "eastAsia"):
                value = fonts.get(qn(f"w:{attr}"))
                if value:
                    values.add(value)
        if allowed_fonts and values and not values.issubset(allowed_fonts):
            findings.append({"rule": "style_font", "style": style_name, "fonts": sorted(values), "severity": "high"})

    for t_idx, table in enumerate(doc.tables, 1):
        pr = table._tbl.tblPr
        width = pr.find(qn("w:tblW"))
        grid = [int(x.get(qn("w:w"), "0")) for x in table._tbl.tblGrid]
        declared = int(width.get(qn("w:w"), "0")) if width is not None and width.get(qn("w:type")) == "dxa" else 0
        if not declared or not grid or sum(grid) != declared:
            findings.append({"rule": "table_geometry", "table": t_idx, "declared": declared, "grid_sum": sum(grid), "severity": "blocker"})

    text = "\n".join(p.text for p in doc.paragraphs)
    for token in profile.get("forbidden_tokens", []):
        if token in text:
            findings.append({"rule": "forbidden_token", "token": token, "severity": "blocker"})

    blockers = [f for f in findings if f["severity"] == "blocker"]
    report = {
        "status": "BLOCKED" if blockers else "REVIEW" if findings else "PASS",
        "sections": len(doc.sections),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "findings": findings,
    }
    payload = json.dumps(report, ensure_ascii=False, indent=2)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(payload, encoding="utf-8")
    print(payload)


if __name__ == "__main__":
    main()
