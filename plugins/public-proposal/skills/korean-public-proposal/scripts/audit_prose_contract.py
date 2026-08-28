#!/usr/bin/env python3
"""Mode-aware Korean public-document prose audit.

The audit detects form and density risks. It does not rewrite text or certify
facts, research validity, human approval, or submission readiness.
"""

from __future__ import annotations

import argparse
import json
import re
import xml.etree.ElementTree as ET
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any


PROFILES = {
    "public_bullet",
    "public_plan",
    "press_release",
    "evaluator_proposal",
    "research_analytic",
    "partnership_brief",
    "executive_brief",
    "official_form_locked",
}
COMPLETE_SENTENCE_END = re.compile(r"(?:다|습니다|됩니다|있습니다|없습니다)[.!?\s]*$")
ANALYTIC_END = re.compile(r"(?:다|습니다|됩니다|있습니다|없습니다|함|있음|없음|됨|임)[.!?\s]*$")
DEONTIC_END = re.compile(r"(?:해야 한다|하여야 한다|해야 함|하여야 함|해야 할 것)[.!?\s]*$")
RHETORIC_CONTRAST = re.compile(
    r"(?:가|이|은|는) 아니라\s*(사람|마음|본질|전환|태도|문화|철학|가치|정신|관계|신뢰|과정|미래|시작|변화|연결)"
)
EVIDENCE_TOKEN = re.compile(
    r"(?:\d+(?:[,.]\d+)*(?:%|명|건|원|백만원|일|개월|년)|\[[A-Z][A-Z0-9_-]*\]|(?:출처|근거|자료|표|그림|붙임|별첨)\s*[:：]?)"
)
PAREN_LABEL = re.compile(r"^[❍○\-–—]?\s*(?:\*\*)?\(([^)]+)\)")


def extract(path: Path) -> str:
    if path.suffix.lower() == ".hwpx":
        return extract_hwpx(path)
    if path.suffix.lower() == ".docx":
        from docx import Document

        document = Document(path)
        parts = [docx_paragraph_text(paragraph) for paragraph in document.paragraphs]
        parts.extend(
            f"| {paragraph.text}"
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            if paragraph.text.strip()
        )
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


def extract_hwpx(path: Path) -> str:
    paragraphs: list[str] = []
    with zipfile.ZipFile(path) as package:
        sections = sorted(
            name for name in package.namelist()
            if re.fullmatch(r"Contents/section\d+\.xml", name)
        )
        for section in sections:
            root = ET.fromstring(package.read(section))
            for paragraph in root.findall("./{http://www.hancom.co.kr/hwpml/2011/paragraph}p"):
                text = hwpx_paragraph_text(paragraph)
                if text:
                    paragraphs.append(normalize_hwpx_line(text))
        if not paragraphs and "Preview/PrvText.txt" in package.namelist():
            preview = package.read("Preview/PrvText.txt").decode("utf-8", errors="replace")
            paragraphs.extend(normalize_hwpx_line(line) for line in preview.splitlines() if line.strip())
    return "\n".join(paragraphs)


def hwpx_paragraph_text(paragraph: ET.Element) -> str:
    skipped = {"tbl", "pic", "container", "rect", "ellipse", "line", "drawText", "ctrl", "secPr"}

    def collect(node: ET.Element) -> list[str]:
        local_name = node.tag.rsplit("}", 1)[-1]
        if local_name in skipped:
            return []
        values = [node.text] if local_name == "t" and node.text else []
        for child in node:
            values.extend(collect(child))
        return values

    return re.sub(r"\s+", " ", "".join(collect(paragraph))).strip()


def normalize_hwpx_line(value: str) -> str:
    value = re.sub(r"^\s*ㅇ\s*", "❍ ", value)
    value = re.sub(r"^\s*[•●]\s*", "❍ ", value)
    return value.strip()


def docx_paragraph_text(paragraph: Any) -> str:
    text = paragraph.text
    style_name = (paragraph.style.name if paragraph.style is not None else "").lower()
    if text.strip() and (
        style_name.startswith("heading")
        or "chapter" in style_name
        or "korean title" in style_name
    ):
        return f"# {text}"
    if text.strip() and "caption" in style_name:
        return f"![caption] {text}"
    if text.strip() and "lead" in style_name:
        return f"> {text}"
    if text.strip() and "note" in style_name:
        return f"※ {text}"
    properties = paragraph._p.pPr
    if text.strip() and properties is not None and properties.numPr is not None:
        return f"- {text}"
    return text


def strip_markup(value: str) -> str:
    value = re.sub(r"^#{1,6}\s+", "", value.strip())
    value = re.sub(r"[*_]{1,2}", "", value)
    return re.sub(r"\s+", " ", value).strip()


def classify(raw: str) -> tuple[str | None, str]:
    stripped = raw.strip()
    indentation = len(raw) - len(raw.lstrip())
    if not stripped or stripped == "---" or stripped.startswith("|") or stripped.startswith("!["):
        return None, ""
    if stripped.startswith("#"):
        return "heading", strip_markup(stripped)
    if stripped.startswith("□"):
        return "square", strip_markup(stripped[1:])
    if stripped.startswith(">"):
        quoted = strip_markup(stripped[1:])
        if re.match(r"^\[(?:도식|그림|표)\s*\d+\]", quoted):
            return "caption", quoted
        return "lead", quoted
    if re.match(r"^(?:⇒|=>)", stripped):
        return "conclusion", strip_markup(re.sub(r"^(?:⇒|=>)\s*", "", stripped))
    if stripped.startswith(("※", "* ")):
        return "note", strip_markup(stripped[1:])
    if re.match(r"^[❍○]\s*", stripped):
        return "item", strip_markup(re.sub(r"^[❍○]\s*", "", stripped))
    if re.match(r"^[-–]\s+", stripped):
        kind = "detail" if indentation >= 2 else "item"
        return kind, strip_markup(re.sub(r"^[-–]\s+", "", stripped))
    return "paragraph", strip_markup(stripped)


def audit(text: str, profile: str, protected: bool = False) -> dict[str, Any]:
    if profile not in PROFILES:
        raise ValueError(f"unknown prose profile: {profile}")

    findings: list[dict[str, Any]] = []
    counts: Counter[str] = Counter()
    metrics: Counter[str] = Counter(
        {
            "headingCount": 0,
            "leadCount": 0,
            "captionCount": 0,
            "noteCount": 0,
            "itemCount": 0,
            "detailCount": 0,
            "conclusionCount": 0,
            "paragraphCount": 0,
            "characters": 0,
            "evidenceBearingUnits": 0,
            "narrativeEndingUnits": 0,
            "completeSentenceParagraphs": 0,
            "analyticParagraphs": 0,
            "completeAnalyticParagraphs": 0,
        }
    )
    lines = text.splitlines()
    in_frontmatter = False

    def add(line: int, severity: str, rule: str, message: str, excerpt: str) -> None:
        findings.append(
            {
                "line": line,
                "severity": severity,
                "rule": rule,
                "message": message,
                "excerpt": excerpt[:160],
            }
        )
        counts[severity] += 1

    if profile == "official_form_locked" and protected:
        return {
            "schemaVersion": "1.0.0",
            "profile": profile,
            "status": "NOT_APPLICABLE",
            "humanReviewRequired": True,
            "scope": "issuer-protected text excluded; audit newly authored free-text fields separately",
            "metrics": {},
            "counts": {},
            "findings": [],
        }

    lead_buffer: list[tuple[int, str]] = []
    current_heading = ""
    for line_number, raw in enumerate(lines, 1):
        if raw.strip() == "---" and (line_number == 1 or in_frontmatter):
            in_frontmatter = not in_frontmatter
            continue
        if in_frontmatter:
            continue
        kind, body = classify(raw)
        if kind is None:
            continue
        if kind == "square":
            if profile == "press_release":
                kind = "item"
            else:
                current_heading = body
                metrics["headingCount"] += 1
                continue
        if kind == "heading":
            current_heading = body
            metrics["headingCount"] += 1
            continue
        if kind not in {"heading", "lead", "caption"} and re.search(
            r"(?:근거|출처|참고문헌|참고자료|sources|references)", current_heading, re.I
        ):
            kind = "note"
        metrics[f"{kind}Count"] += 1
        metrics["characters"] += len(body)
        if EVIDENCE_TOKEN.search(body):
            metrics["evidenceBearingUnits"] += 1
        if COMPLETE_SENTENCE_END.search(body):
            metrics["narrativeEndingUnits"] += 1
        if kind == "paragraph" and ANALYTIC_END.search(body):
            metrics["completeSentenceParagraphs"] += 1
        if kind == "paragraph" and len(body) >= 40:
            metrics["analyticParagraphs"] += 1
            if ANALYTIC_END.search(body):
                metrics["completeAnalyticParagraphs"] += 1

        if re.search(r"[!！]", body) and kind != "note":
            if profile in {"public_bullet", "public_plan", "partnership_brief", "executive_brief"}:
                add(line_number, "error", "EXCLAMATION", "본문의 감탄 표현은 공공문서 기본 문법에서 벗어남", body)
            else:
                add(line_number, "info", "EXCLAMATION_CONTEXT", "인용 제목·자료값인지 확인; 홍보성 감탄이면 수정", body)
        if re.search(r"[?？]", body) and kind != "note":
            if profile in {"public_bullet", "public_plan", "partnership_brief", "executive_brief"}:
                add(line_number, "error", "RHETORICAL_QUESTION", "개조식 본문의 질문 표현을 판단·요청·확인항목으로 수정", body)
            else:
                add(line_number, "info", "QUESTION_CONTEXT", "연구질문·평가척도·RFP 인용인지 확인; 수사적 질문이면 수정", body)
        if re.search(r"(?:것이다|것임)\b", body):
            add(line_number, "warning", "GEOSIDA", "'것이다/것임'을 구체 판단이나 명사로 바꿀 수 있는지 검토", body)
        if RHETORIC_CONTRAST.search(body):
            if profile in {"public_bullet", "public_plan", "partnership_brief", "executive_brief"}:
                add(line_number, "error", "RHETORIC_CONTRAST", "추상어를 이용한 'A가 아니라 B' 수사", body)
            else:
                add(
                    line_number,
                    "info",
                    "CONTRAST_CONTEXT",
                    "책임·범위·대안의 실질적 경계인지 확인; 추상적 강조에 그치면 직접 판단문으로 수정",
                    body,
                )
        if "—" in body and kind in {"item", "detail", "conclusion"}:
            add(line_number, "warning", "EM_DASH_SLOGAN", "줄표로 이어 붙인 문구를 별도 근거 또는 세부 항목으로 분리할지 검토", body)
        label = PAREN_LABEL.search(raw.strip()) if kind in {"item", "detail"} else None
        if label and not 2 <= len(re.sub(r"\s+", "", label.group(1))) <= 14:
            add(line_number, "warning", "PAREN_LABEL_LENGTH", "괄호형 선행 표제는 짧고 구체적으로 작성", body)

        if profile in {"public_bullet", "public_plan", "partnership_brief", "executive_brief"}:
            if kind in {"item", "detail", "conclusion"} and COMPLETE_SENTENCE_END.search(body):
                severity = "error" if profile == "public_bullet" else "warning"
                add(line_number, severity, "BULLET_NARRATIVE_ENDING", "개조식 항목은 명사구 또는 필요한 경우 '~함/~있음' 종결 권장", body)
            item_limit = 90 if profile == "public_plan" else 70
            detail_limit = 120 if profile == "public_plan" else 100
            if kind == "item" and len(body) > item_limit:
                add(line_number, "warning", "ITEM_LONG", f"핵심 항목이 {item_limit}자를 초과함; 근거를 세부 항목으로 분리할지 검토", body)
            if kind == "detail" and len(body) > detail_limit:
                add(line_number, "warning", "DETAIL_LONG", f"세부 항목이 {detail_limit}자를 초과함; 한 줄 한 생각 원칙 점검", body)
            if kind == "conclusion" and len(body) > 60:
                add(line_number, "warning", "CONCLUSION_LONG", "결론이 60자를 초과함; 한 판단만 남길지 검토", body)
            if kind == "conclusion" and DEONTIC_END.search(body):
                add(line_number, "error", "DEONTIC_CONCLUSION", "결론의 당위 표현을 근거에서 도출된 구체 판단으로 수정", body)

        if profile == "press_release" and kind == "item" and len(body) > 170:
            add(line_number, "warning", "PRESS_ITEM_LONG", "보도자료 항목이 170자를 초과함; 발표 사실·배경·인용을 분리할지 검토", body)

        if kind == "lead":
            lead_buffer.append((line_number, body))
            next_line = lines[line_number] if line_number < len(lines) else ""
            if not next_line.strip().startswith(">"):
                lead = " ".join(part for _, part in lead_buffer)
                first_line = lead_buffer[0][0]
                lead_buffer.clear()
                if profile == "public_bullet" and not re.search(r"(?:고자 함|하려 함|려고 함)\.?$", lead):
                    add(first_line, "warning", "LEAD_ENDING", "개조식 리드문은 한 문장 '~하고자 함.' 형식을 우선 검토", lead)
                if len(lead) > 140:
                    add(first_line, "warning", "LEAD_LONG", "리드문이 140자를 초과함", lead)

        if kind == "paragraph":
            if profile == "evaluator_proposal" and len(body) > 650:
                add(line_number, "warning", "EVALUATOR_PARAGRAPH_LONG", "평가 대응 문단이 과도하게 길어 직접 답변·근거·통제로 분리할지 검토", body)
            if profile == "research_analytic" and len(body) > 900:
                add(line_number, "warning", "RESEARCH_PARAGRAPH_LONG", "연구 문단이 과도하게 길어 관찰·분석·추론·한계로 분리할지 검토", body)

    substantive_units = sum(metrics[f"{kind}Count"] for kind in ("item", "detail", "conclusion", "paragraph"))
    metrics["substantiveUnits"] = substantive_units
    metrics["narrativeEndingRatioPermille"] = round(
        metrics["narrativeEndingUnits"] * 1000 / substantive_units
    ) if substantive_units else 0
    metrics["evidenceBearingRatioPermille"] = round(
        metrics["evidenceBearingUnits"] * 1000 / substantive_units
    ) if substantive_units else 0
    metrics["completeSentenceParagraphRatioPermille"] = round(
        metrics["completeSentenceParagraphs"] * 1000 / metrics["paragraphCount"]
    ) if metrics["paragraphCount"] else 0
    metrics["completeAnalyticParagraphRatioPermille"] = round(
        metrics["completeAnalyticParagraphs"] * 1000 / metrics["analyticParagraphs"]
    ) if metrics["analyticParagraphs"] else 0

    if profile in {"evaluator_proposal", "research_analytic"} and metrics["characters"] >= 1200:
        if metrics["evidenceBearingUnits"] == 0:
            add(0, "warning", "NO_VISIBLE_EVIDENCE", "장문 원고에 수치·출처·근거·표·그림 로케이터가 보이지 않음", "")
        if metrics["paragraphCount"] == 0:
            add(0, "warning", "NO_ANALYTIC_PARAGRAPHS", "분석 또는 평가 대응 원고가 표·개조식만으로 구성됨", "")
        elif metrics["analyticParagraphs"] >= 3 and metrics["completeAnalyticParagraphRatioPermille"] < 600:
            add(0, "warning", "FRAGMENT_HEAVY_ANALYSIS", "분석 문단의 완결문 비율이 낮음; 표제·명사구가 논증을 대체하는지 검토", "")

    status = "BLOCKED" if counts["error"] else "REVIEW" if counts["warning"] else "PASS"
    return {
        "schemaVersion": "1.0.0",
        "profile": profile,
        "status": status,
        "humanReviewRequired": True,
        "scope": "form and density risk only",
        "metrics": dict(metrics),
        "counts": dict(counts),
        "findings": findings,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Audit a mode-aware Korean prose contract.")
    parser.add_argument("input", type=Path)
    parser.add_argument("--profile", choices=sorted(PROFILES), required=True)
    parser.add_argument("--protected", action="store_true", help="Treat the input as issuer-protected text.")
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()
    report = audit(extract(args.input), args.profile, protected=args.protected)
    payload = json.dumps(report, ensure_ascii=False, indent=2)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(payload + "\n", encoding="utf-8")
    print(payload)
    return 2 if report["status"] == "BLOCKED" else 0


if __name__ == "__main__":
    raise SystemExit(main())
