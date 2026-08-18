#!/usr/bin/env python3
"""Detect common AI-proposal patterns without rewriting source text."""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path


BUZZWORDS = (
    "혁신적",
    "획기적",
    "선도적",
    "압도적",
    "체계적",
    "실효적",
    "유기적",
    "최적의",
    "최고의",
    "완벽한",
)
SIGNATURES = (
    "이를 통해",
    "결론적으로",
    "다음과 같은",
    "시사하는 바가 크다",
    "할 수 있다",
    "할 예정이다",
    "고려한다",
)
PLACEHOLDERS = ("TODO", "TBD", "None", "placeholder", "연결된 Figure", "[확인 필요]")


def extract(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        from docx import Document

        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        parts += [p.text for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs]
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


def lint(text: str) -> dict:
    findings = []
    for term in BUZZWORDS:
        count = text.count(term)
        if count >= 3:
            findings.append({"rule": "hype_repetition", "term": term, "count": count, "severity": "high"})
    for term in SIGNATURES:
        count = text.count(term)
        threshold = 3 if term != "할 수 있다" else 5
        if count >= threshold:
            findings.append({"rule": "ai_signature", "term": term, "count": count, "severity": "medium"})
    for term in PLACEHOLDERS:
        count = text.count(term)
        if count:
            findings.append({"rule": "unresolved_placeholder", "term": term, "count": count, "severity": "blocker"})
    normalized = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    repeated = Counter(line for line in normalized if len(line) >= 28)
    for line, count in repeated.most_common(20):
        if count >= 3:
            findings.append({"rule": "verbatim_repetition", "text": line[:160], "count": count, "severity": "high"})
    colon_heads = sum(1 for line in normalized if re.match(r"^(?:\d+[.)]?\s*)?[^:]{2,28}:\s+", line))
    if colon_heads >= 6:
        findings.append({"rule": "colon_heading_formula", "count": colon_heads, "severity": "medium"})
    future_modals = len(re.findall(r"할\s+수\s+있(?:다|습니다)|할\s+예정", text))
    if future_modals >= 8:
        findings.append({"rule": "noncommittal_future", "count": future_modals, "severity": "high"})
    return {
        "characters": len(text),
        "findings": findings,
        "counts": Counter(item["severity"] for item in findings),
        "status": "BLOCKED" if any(x["severity"] == "blocker" for x in findings) else "REVIEW" if findings else "PASS",
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()
    report = lint(extract(args.input))
    payload = json.dumps(report, ensure_ascii=False, indent=2, default=dict)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(payload, encoding="utf-8")
    print(payload)


if __name__ == "__main__":
    main()
