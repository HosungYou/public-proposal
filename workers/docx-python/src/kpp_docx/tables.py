"""Fixed-DXA Word-native table construction."""

from __future__ import annotations

from dataclasses import dataclass

from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell

from .styles import TypographyContract, format_body_paragraph, format_run


@dataclass(frozen=True)
class TableContract:
    width_dxa: int
    margin_top_dxa: int
    margin_start_dxa: int
    margin_bottom_dxa: int
    margin_end_dxa: int
    border_size_eighth_pt: int


def add_native_table(
    document: DocumentType,
    *,
    headers: list[str],
    rows: list[list[str]],
    column_widths_dxa: list[int],
    contract: TableContract,
    typography: TypographyContract,
) -> Table:
    """Create a native fixed-layout table with explicit grid and cell geometry."""

    if not headers:
        raise ValueError("table headers must not be empty")
    if len(column_widths_dxa) != len(headers):
        raise ValueError("table column widths must match header count")
    if sum(column_widths_dxa) != contract.width_dxa:
        raise ValueError("table column widths must sum to the locked table width")
    if any(len(row) != len(headers) for row in rows):
        raise ValueError("every table row must match the header count")

    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    _apply_table_properties(table, contract, column_widths_dxa)
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(repeat_header)

    for index, value in enumerate(headers):
        _set_cell_text(
            table.rows[0].cells[index],
            value,
            width_dxa=column_widths_dxa[index],
            typography=typography,
            header=True,
        )
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            _set_cell_text(
                cells[index],
                value,
                width_dxa=column_widths_dxa[index],
                typography=typography,
                header=False,
            )
    return table


def _apply_table_properties(
    table: Table,
    contract: TableContract,
    column_widths_dxa: list[int],
) -> None:
    properties = table._tbl.tblPr
    for local_name in ("tblStyle", "tblW", "tblLayout", "tblCellMar", "tblBorders"):
        existing = properties.find(qn(f"w:{local_name}"))
        if existing is not None:
            properties.remove(existing)

    width = OxmlElement("w:tblW")
    width.set(qn("w:w"), str(contract.width_dxa))
    width.set(qn("w:type"), "dxa")
    properties.append(width)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    properties.append(layout)

    margins = OxmlElement("w:tblCellMar")
    for side, value in (
        ("top", contract.margin_top_dxa),
        ("start", contract.margin_start_dxa),
        ("bottom", contract.margin_bottom_dxa),
        ("end", contract.margin_end_dxa),
    ):
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    properties.append(margins)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(contract.border_size_eighth_pt))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "7F8C99")
        borders.append(element)
    properties.append(borders)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in column_widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)


def _set_cell_text(
    cell: _Cell,
    value: str,
    *,
    width_dxa: int,
    typography: TypographyContract,
    header: bool,
) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    existing_width = cell_properties.find(qn("w:tcW"))
    if existing_width is not None:
        cell_properties.remove(existing_width)
    width = OxmlElement("w:tcW")
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")
    cell_properties.insert(0, width)

    existing_shading = cell_properties.find(qn("w:shd"))
    if existing_shading is not None:
        cell_properties.remove(existing_shading)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "E8EEF5" if header else "FFFFFF")
    cell_properties.append(shading)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(value)
    if header:
        paragraph.style = "KPP Table Header"
        format_run(
            run,
            font=typography.heading_font,
            half_points=18,
            bold=True,
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        format_body_paragraph(paragraph, typography)
        # Preserve the governed body rhythm without reclassifying table prose as body copy.
        paragraph.style = "KPP Table Body"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
