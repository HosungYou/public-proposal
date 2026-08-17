"""Governed Word-native proposal document builder."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Twips
from pydantic import BaseModel, ConfigDict, Field, model_validator

from .styles import (
    TypographyContract,
    format_body_paragraph,
    format_navigation_paragraph,
    format_run,
    install_governed_styles,
)
from .tables import TableContract, add_native_table


SCHEMA_VERSION = "1.0.0"
BUILDER_VERSION = "0.1.0"
SHA256_PATTERN = r"^[a-f0-9]{64}$"


class StrictModel(BaseModel):
    model_config = ConfigDict(extra="forbid", populate_by_name=True)


class TemplateRef(StrictModel):
    asset_id: str = Field(alias="assetId", min_length=1)
    path: str = Field(min_length=1)
    sha256: str = Field(pattern=SHA256_PATTERN)


class FigureSpec(StrictModel):
    figure_id: str = Field(alias="figureId", min_length=1)
    page_id: str = Field(alias="pageId", min_length=1)
    path: str = Field(min_length=1)
    sha256: str = Field(pattern=SHA256_PATTERN)
    format: Literal["png", "jpeg", "jpg"]
    caption: str = Field(min_length=1)
    evidence_ids: list[str] = Field(alias="evidenceIds")
    width_dxa: int = Field(alias="widthDxa", gt=0, default=7200)


class FigureManifest(StrictModel):
    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    figures: list[FigureSpec]


class PagePlanItem(StrictModel):
    page_id: str = Field(alias="pageId", min_length=1)
    requirement_id: str = Field(alias="requirementId", min_length=1)
    page_role: str = Field(alias="pageRole", min_length=1)
    surface_template_id: str = Field(alias="surfaceTemplateId", min_length=1)
    claim_ids: list[str] = Field(alias="claimIds")
    figure_specs: list[dict[str, object]] = Field(alias="figureSpecs")


class PagePlan(StrictModel):
    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    pages: list[PagePlanItem] = Field(min_length=1)


class EvidenceClaim(StrictModel):
    claim_id: str = Field(alias="claimId", min_length=1)
    status: Literal["verified", "bounded", "pending_blank", "blocked"]
    evidence_ids: list[str] = Field(alias="evidenceIds")


class EvidenceBinding(StrictModel):
    evidence_id: str = Field(alias="evidenceId", min_length=1)
    source_path: str = Field(alias="sourcePath", min_length=1)
    source_sha256: str = Field(alias="sourceSha256", pattern=SHA256_PATTERN)
    scope: str = Field(min_length=1)
    claim_ids: list[str] = Field(alias="claimIds", min_length=1)
    target_requirement_id: str = Field(alias="targetRequirementId", min_length=1)
    target_page_id: str = Field(alias="targetPageId", min_length=1)
    target_page_role: str = Field(alias="targetPageRole", min_length=1)


class EvidenceLedger(StrictModel):
    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    claims: list[EvidenceClaim]
    bindings: list[EvidenceBinding]


class ParagraphContent(StrictModel):
    text: str
    claim_ids: list[str] = Field(alias="claimIds")
    evidence_ids: list[str] = Field(alias="evidenceIds")


class TableContent(StrictModel):
    table_id: str = Field(alias="tableId", min_length=1)
    caption: str = Field(min_length=1)
    headers: list[str] = Field(min_length=1)
    rows: list[list[str]]
    column_widths_dxa: list[int] = Field(alias="columnWidthsDxa", min_length=1)


class ContentBlock(StrictModel):
    page_id: str = Field(alias="pageId", min_length=1)
    heading: str = Field(min_length=1)
    paragraphs: list[ParagraphContent]
    tables: list[TableContent]
    figure_ids: list[str] = Field(alias="figureIds")


class TypographyProfile(StrictModel):
    heading_font: Literal["Noto Sans CJK KR"] = Field(alias="headingFont")
    navigation_font: Literal["Noto Sans CJK KR"] = Field(alias="navigationFont")
    body_font: Literal["Noto Serif CJK KR"] = Field(alias="bodyFont")
    body_point: Literal[9.3] = Field(alias="bodyPoint")
    line_height: Literal[1.52] = Field(alias="lineHeight")
    alignment: Literal["justified"]
    character_spacing_pt: float = Field(alias="characterSpacingPt", ge=-2, le=2)


class CellMargins(StrictModel):
    top: int = Field(ge=0)
    start: int = Field(ge=0)
    bottom: int = Field(ge=0)
    end: int = Field(ge=0)


class TableProfile(StrictModel):
    width_dxa: int = Field(alias="widthDxa", gt=0)
    cell_margin_dxa: CellMargins = Field(alias="cellMarginDxa")
    border_size_eighth_pt: int = Field(alias="borderSizeEighthPt", gt=0)


class SurfaceProfile(StrictModel):
    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    profile_id: str = Field(alias="profileId", min_length=1)
    status: Literal["locked"]
    typography: TypographyProfile
    table: TableProfile


class BuildOutput(StrictModel):
    docx_path: str = Field(alias="docxPath", min_length=1)
    manifest_path: str = Field(alias="manifestPath", min_length=1)


class BuildRequest(StrictModel):
    """Versioned, closed input boundary for one document build."""

    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    project_id: str = Field(alias="projectId", min_length=1)
    template: TemplateRef
    page_plan: PagePlan = Field(alias="pagePlan")
    evidence_ledger: EvidenceLedger = Field(alias="evidenceLedger")
    content_blocks: list[ContentBlock] = Field(alias="contentBlocks", min_length=1)
    figure_manifest: FigureManifest = Field(alias="figureManifest")
    surface_profile: SurfaceProfile = Field(alias="surfaceProfile")
    output: BuildOutput

    @model_validator(mode="after")
    def validate_cross_references(self) -> "BuildRequest":
        page_ids = [page.page_id for page in self.page_plan.pages]
        if len(page_ids) != len(set(page_ids)):
            raise ValueError("pagePlan pageId values must be unique")
        content_page_ids = [block.page_id for block in self.content_blocks]
        if len(content_page_ids) != len(set(content_page_ids)):
            raise ValueError("content block pageId values must be unique")
        if set(content_page_ids) != set(page_ids):
            raise ValueError("content blocks must map exactly to pagePlan pages")

        figures = {figure.figure_id: figure for figure in self.figure_manifest.figures}
        if len(figures) != len(self.figure_manifest.figures):
            raise ValueError("figure manifest figureId values must be unique")
        referenced_figures = {
            figure_id for block in self.content_blocks for figure_id in block.figure_ids
        }
        if referenced_figures != set(figures):
            raise ValueError("content figureIds must map exactly to the figure manifest")
        if any(figure.page_id not in page_ids for figure in figures.values()):
            raise ValueError("figure pageId must reference a planned page")

        evidence_ids = [binding.evidence_id for binding in self.evidence_ledger.bindings]
        if len(evidence_ids) != len(set(evidence_ids)):
            raise ValueError("evidence ledger evidenceId values must be unique")
        claim_ids = [claim.claim_id for claim in self.evidence_ledger.claims]
        if len(claim_ids) != len(set(claim_ids)):
            raise ValueError("evidence ledger claimId values must be unique")
        known_evidence_ids = set(evidence_ids)
        known_claim_ids = set(claim_ids)

        claim_by_id = {claim.claim_id: claim for claim in self.evidence_ledger.claims}
        binding_by_id = {
            binding.evidence_id: binding for binding in self.evidence_ledger.bindings
        }
        for claim in self.evidence_ledger.claims:
            if not set(claim.evidence_ids).issubset(known_evidence_ids):
                raise ValueError("claim evidenceIds must exist in the evidence ledger")
            if any(
                claim.claim_id not in binding_by_id[evidence_id].claim_ids
                for evidence_id in claim.evidence_ids
            ):
                raise ValueError("claim and evidence binding links must be reciprocal")
        for binding in self.evidence_ledger.bindings:
            if not set(binding.claim_ids).issubset(known_claim_ids):
                raise ValueError("evidence binding claimIds must exist in the evidence ledger")
            if any(
                binding.evidence_id not in claim_by_id[claim_id].evidence_ids
                for claim_id in binding.claim_ids
            ):
                raise ValueError("claim and evidence binding links must be reciprocal")

        page_by_id = {page.page_id: page for page in self.page_plan.pages}
        for page in self.page_plan.pages:
            if not set(page.claim_ids).issubset(known_claim_ids):
                raise ValueError("page claimIds must exist in the evidence ledger")
        for binding in self.evidence_ledger.bindings:
            target_page = page_by_id.get(binding.target_page_id)
            if target_page is None:
                raise ValueError("evidence binding targetPageId must reference a planned page")
            if binding.target_requirement_id != target_page.requirement_id:
                raise ValueError("evidence binding requirement must match its planned page")
            if binding.target_page_role != target_page.page_role:
                raise ValueError("evidence binding page role must match its planned page")
            if not set(binding.claim_ids).issubset(set(target_page.claim_ids)):
                raise ValueError("evidence binding claimIds must be declared by its target page")

        for block in self.content_blocks:
            planned_claim_ids = set(page_by_id[block.page_id].claim_ids)
            for paragraph in block.paragraphs:
                if not set(paragraph.claim_ids).issubset(planned_claim_ids):
                    raise ValueError("paragraph claimIds must be declared by its planned page")
                if not set(paragraph.evidence_ids).issubset(known_evidence_ids):
                    raise ValueError("paragraph evidenceIds must exist in the evidence ledger")
                if any(
                    binding_by_id[evidence_id].target_page_id != block.page_id
                    for evidence_id in paragraph.evidence_ids
                ):
                    raise ValueError("paragraph evidenceIds must target its planned page")
            for figure_id in block.figure_ids:
                figure = figures[figure_id]
                if figure.page_id != block.page_id:
                    raise ValueError("figure pageId must match the content block pageId")
                if not set(figure.evidence_ids).issubset(known_evidence_ids):
                    raise ValueError("figure evidenceIds must exist in the evidence ledger")
                if any(
                    binding_by_id[evidence_id].target_page_id != block.page_id
                    for evidence_id in figure.evidence_ids
                ):
                    raise ValueError("figure evidenceIds must target its planned page")

        table_width = self.surface_profile.table.width_dxa
        for block in self.content_blocks:
            for table in block.tables:
                if len(table.column_widths_dxa) != len(table.headers):
                    raise ValueError("table widths must match the header count")
                if sum(table.column_widths_dxa) != table_width:
                    raise ValueError("table widths must sum to the locked profile width")
                if any(len(row) != len(table.headers) for row in table.rows):
                    raise ValueError("table rows must match the header count")
        return self


class BuildResult(BaseModel):
    model_config = ConfigDict(arbitrary_types_allowed=True)
    docx: Path
    manifest: Path


def build_document(request: BuildRequest) -> BuildResult:
    """Build a proposal DOCX and a manifest bound to its exact inputs and output."""

    template_path = Path(request.template.path).expanduser().resolve()
    if not template_path.is_file():
        raise ValueError(f"template is missing: {template_path}")
    template_hash = _sha256_file(template_path)
    if template_hash != request.template.sha256:
        raise ValueError("template SHA-256 does not match the locked BuildRequest")

    figure_records: list[dict[str, object]] = []
    for figure in request.figure_manifest.figures:
        figure_path = Path(figure.path).expanduser().resolve()
        if not figure_path.is_file():
            raise ValueError(f"figure is missing: {figure_path}")
        if _sha256_file(figure_path) != figure.sha256:
            raise ValueError(f"figure SHA-256 does not match: {figure.figure_id}")
        figure_records.append(
            {
                "figureId": figure.figure_id,
                "pageId": figure.page_id,
                "path": str(figure_path),
                "sha256": figure.sha256,
                "format": figure.format,
                "caption": figure.caption,
                "evidenceIds": figure.evidence_ids,
                "embedded": True,
            }
        )

    docx_path = Path(request.output.docx_path).expanduser().resolve()
    manifest_path = Path(request.output.manifest_path).expanduser().resolve()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document(template_path)
    _clear_document_body(document)

    typography = _typography_contract(request.surface_profile.typography)
    style_ids = install_governed_styles(document, typography)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            format_navigation_paragraph(paragraph, typography)
        for paragraph in section.footer.paragraphs:
            format_navigation_paragraph(paragraph, typography)
    table_contract = _table_contract(request.surface_profile.table)
    content_by_page_id = {block.page_id: block for block in request.content_blocks}
    figure_by_id = {figure.figure_id: figure for figure in request.figure_manifest.figures}
    table_records: list[dict[str, object]] = []

    for block_index, planned_page in enumerate(request.page_plan.pages):
        block = content_by_page_id[planned_page.page_id]
        if block_index:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        heading = document.add_paragraph(style=style_ids["heading"])
        heading_run = heading.add_run(block.heading)
        format_run(
            heading_run,
            font=typography.heading_font,
            half_points=32,
            bold=True,
        )
        for paragraph_content in block.paragraphs:
            paragraph = document.add_paragraph()
            paragraph.add_run(paragraph_content.text)
            format_body_paragraph(paragraph, typography)

        for table_content in block.tables:
            caption = document.add_paragraph(style=style_ids["caption"])
            caption_run = caption.add_run(table_content.caption)
            format_run(
                caption_run,
                font=typography.heading_font,
                half_points=18,
                bold=True,
            )
            add_native_table(
                document,
                headers=table_content.headers,
                rows=table_content.rows,
                column_widths_dxa=table_content.column_widths_dxa,
                contract=table_contract,
                typography=typography,
            )
            table_records.append(
                {
                    "tableId": table_content.table_id,
                    "pageId": block.page_id,
                    "widthDxa": table_contract.width_dxa,
                    "columnWidthsDxa": table_content.column_widths_dxa,
                    "native": True,
                }
            )

        for figure_id in block.figure_ids:
            figure = figure_by_id[figure_id]
            caption = document.add_paragraph(style=style_ids["caption"])
            caption_run = caption.add_run(figure.caption)
            format_run(
                caption_run,
                font=typography.heading_font,
                half_points=18,
                bold=True,
            )
            document.add_picture(
                str(Path(figure.path).expanduser().resolve()),
                width=Twips(figure.width_dxa),
            )

    document.save(docx_path)
    docx_hash = _sha256_file(docx_path)

    manifest = {
        "schemaVersion": SCHEMA_VERSION,
        "builderVersion": BUILDER_VERSION,
        "projectId": request.project_id,
        "template": {
            "assetId": request.template.asset_id,
            "path": str(template_path),
            "sha256": template_hash,
        },
        "profile": {
            "profileId": request.surface_profile.profile_id,
            "status": request.surface_profile.status,
            "sha256": _sha256_json(request.surface_profile),
        },
        "inputs": {
            "pagePlanSha256": _sha256_json(request.page_plan),
            "evidenceLedgerSha256": _sha256_json(request.evidence_ledger),
            "contentBlocksSha256": _sha256_json(request.content_blocks),
            "figureManifestSha256": _sha256_json(request.figure_manifest),
            "surfaceProfileSha256": _sha256_json(request.surface_profile),
        },
        "pages": [
            {
                "pageId": page.page_id,
                "pageRole": page.page_role,
                "surfaceTemplateId": page.surface_template_id,
            }
            for page in request.page_plan.pages
        ],
        "styles": {
            "heading": {"font": typography.heading_font},
            "navigation": {"font": typography.navigation_font},
            "body": {
                "font": typography.body_font,
                "point": typography.body_point,
                "ooxmlHalfPoints": typography.body_half_points,
                "effectiveOoxmlPoint": typography.body_half_points / 2,
                "lineHeight": typography.line_height,
                "lineDxa": typography.body_line_dxa,
                "alignment": "justified",
                "characterSpacingPt": typography.character_spacing_pt,
                "characterSpacingTwips": typography.character_spacing_twips,
            },
        },
        "tables": table_records,
        "figures": figure_records,
        "artifacts": {
            "docx": {"path": str(docx_path), "sha256": docx_hash},
        },
    }
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, sort_keys=True, indent=2) + "\n",
        encoding="utf-8",
    )
    return BuildResult(docx=docx_path, manifest=manifest_path)


def _clear_document_body(document: object) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _typography_contract(profile: TypographyProfile) -> TypographyContract:
    return TypographyContract(
        heading_font=profile.heading_font,
        navigation_font=profile.navigation_font,
        body_font=profile.body_font,
        body_point=profile.body_point,
        line_height=profile.line_height,
        character_spacing_pt=profile.character_spacing_pt,
    )


def _table_contract(profile: TableProfile) -> TableContract:
    return TableContract(
        width_dxa=profile.width_dxa,
        margin_top_dxa=profile.cell_margin_dxa.top,
        margin_start_dxa=profile.cell_margin_dxa.start,
        margin_bottom_dxa=profile.cell_margin_dxa.bottom,
        margin_end_dxa=profile.cell_margin_dxa.end,
        border_size_eighth_pt=profile.border_size_eighth_pt,
    )


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _sha256_json(value: object) -> str:
    if isinstance(value, BaseModel):
        value = value.model_dump(by_alias=True)
    elif isinstance(value, list):
        value = [item.model_dump(by_alias=True) if isinstance(item, BaseModel) else item for item in value]
    encoded = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()
