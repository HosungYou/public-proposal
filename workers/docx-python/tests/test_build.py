"""Artifact-level tests for governed Word-native proposal builds."""

from __future__ import annotations

import hashlib
import json
import os
import zipfile
from base64 import b64decode
from pathlib import Path

import pytest
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pydantic import ValidationError

import kpp_docx.build as build_module
from kpp_docx.build import BuildRequest, build_document


WORKER_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = WORKER_ROOT / "assets" / "Korean Public Proposal A4 v1.docx"


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _xml(path: Path, member: str) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read(member).decode("utf-8")


def _canonical_sha(value: object) -> str:
    return hashlib.sha256(
        json.dumps(
            value,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()


def sample_request(tmp_path: Path) -> BuildRequest:
    return BuildRequest.model_validate(
        {
            "schemaVersion": "1.0.0",
            "projectId": "synthetic-research-proposal",
            "template": {
                "assetId": "korean-public-proposal-a4-v1",
                "path": str(TEMPLATE),
                "sha256": _sha256(TEMPLATE),
            },
            "pagePlan": {
                "schemaVersion": "1.0.0",
                "pages": [
                    {
                        "pageId": "P-01",
                        "requirementId": "REQ-01",
                        "pageRole": "research_method",
                        "surfaceTemplateId": "r08-research-method-v1",
                        "claimIds": ["CLM-01"],
                        "figureSpecs": [],
                    }
                ],
            },
            "evidenceLedger": {
                "schemaVersion": "1.0.0",
                "claims": [
                    {
                        "claimId": "CLM-01",
                        "status": "verified",
                        "evidenceIds": ["EV-01"],
                    }
                ],
                "bindings": [
                    {
                        "evidenceId": "EV-01",
                        "sourcePath": "/locked/source.pdf",
                        "sourceSha256": "a" * 64,
                        "scope": "연구 수행방법의 공식 근거",
                        "claimIds": ["CLM-01"],
                        "targetRequirementId": "REQ-01",
                        "targetPageId": "P-01",
                        "targetPageRole": "research_method",
                    }
                ],
            },
            "contentBlocks": [
                {
                    "pageId": "P-01",
                    "heading": "1. 연구 수행방법",
                    "paragraphs": [
                        {
                            "text": "공식 자료와 현장 검증을 연결하여 연구 결과의 활용 가능성을 높인다.",
                            "claimIds": ["CLM-01"],
                            "evidenceIds": ["EV-01"],
                        }
                    ],
                    "tables": [
                        {
                            "tableId": "TBL-01",
                            "caption": "표 1. 연구 단계별 산출물",
                            "headers": ["단계", "산출물"],
                            "rows": [["착수", "연구설계서"], ["분석", "중간보고서"]],
                            "columnWidthsDxa": [2400, 6000],
                        }
                    ],
                    "figureIds": [],
                }
            ],
            "figureManifest": {"schemaVersion": "1.0.0", "figures": []},
            "surfaceProfile": {
                "schemaVersion": "1.0.0",
                "profileId": "synthetic-korean-public-research-v1",
                "status": "locked",
                "typography": {
                    "headingFont": "Noto Sans CJK KR",
                    "navigationFont": "Noto Sans CJK KR",
                    "bodyFont": "Noto Serif CJK KR",
                    "bodyPoint": 9.3,
                    "lineHeight": 1.52,
                    "alignment": "justified",
                    "characterSpacingPt": -0.2,
                    "precisionPolicy": "acknowledged_half_point_quantization",
                },
                "table": {
                    "widthDxa": 8400,
                    "cellMarginDxa": {"top": 80, "start": 100, "bottom": 80, "end": 100},
                    "borderSizeEighthPt": 4,
                },
            },
            "output": {
                "docxPath": str(tmp_path / "proposal.docx"),
                "manifestPath": str(tmp_path / "build-manifest.json"),
            },
        }
    )


def test_build_applies_body_and_table_contract(tmp_path: Path) -> None:
    result = build_document(sample_request(tmp_path))

    document_xml = _xml(result.docx, "word/document.xml")
    styles_xml = _xml(result.docx, "word/styles.xml")
    header_xml = _xml(result.docx, "word/header1.xml")

    assert 'w:jc w:val="both"' in document_xml
    assert 'w:spacing w:line="365" w:lineRule="auto"' in document_xml
    assert 'w:spacing w:val="-4"' in document_xml
    assert 'w:tblLayout w:type="fixed"' in document_xml
    assert 'w:gridCol w:w="2400"' in document_xml
    assert 'w:gridCol w:w="6000"' in document_xml
    assert "w:tblCellMar" in document_xml
    assert "TableGrid" not in document_xml
    assert "Noto Sans CJK KR" in styles_xml
    assert "Noto Serif CJK KR" in styles_xml
    assert 'w:sz w:val="19"' in styles_xml
    assert "Noto Sans CJK KR" in header_xml

    reopened = Document(result.docx)
    assert reopened.paragraphs[0].text == "1. 연구 수행방법"
    assert len(reopened.tables) == 1
    assert reopened.sections[0].header.paragraphs[0].text == "연구용역 제안서 | 형식·구조 검증본"
    assert reopened.sections[0].footer.paragraphs[0].text.startswith("1. 연구 수행방법 |")


def test_native_table_applies_governed_header_and_body_cell_grammar(tmp_path: Path) -> None:
    result = build_document(sample_request(tmp_path))
    table = Document(result.docx).tables[0]

    margins = table._tbl.tblPr.find(qn("w:tblCellMar"))
    assert margins is not None
    assert {
        side: margins.find(qn(f"w:{side}")).get(qn("w:w"))
        for side in ("top", "start", "bottom", "end")
    } == {"top": "80", "start": "100", "bottom": "80", "end": "100"}

    header_row = table.rows[0]
    assert header_row._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
    for cell in header_row.cells:
        shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
        assert shading is not None
        assert shading.get(qn("w:fill")) == "E8EEF5"
        assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in table.rows[1:]:
        for cell in row.cells:
            shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            assert shading is not None
            assert shading.get(qn("w:fill")) == "FFFFFF"
            assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
            assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
            spacing = cell.paragraphs[0]._p.get_or_add_pPr().find(qn("w:spacing"))
            assert spacing is not None
            assert spacing.get(qn("w:line")) == "365"
            assert spacing.get(qn("w:lineRule")) == "auto"


def test_manifest_binds_template_inputs_pages_and_output_hashes(tmp_path: Path) -> None:
    request = sample_request(tmp_path)
    result = build_document(request)
    manifest = json.loads(result.manifest.read_text(encoding="utf-8"))

    assert manifest["schemaVersion"] == "1.0.0"
    assert manifest["projectId"] == request.project_id
    assert manifest["template"] == {
        "assetId": "korean-public-proposal-a4-v1",
        "path": str(TEMPLATE.resolve()),
        "sha256": _sha256(TEMPLATE),
    }
    assert manifest["pages"] == [
        {
            "pageId": "P-01",
            "pageRole": "research_method",
            "surfaceTemplateId": "r08-research-method-v1",
        }
    ]
    assert manifest["styles"]["body"] == {
        "font": "Noto Serif CJK KR",
        "requestedPoint": 9.3,
        "precisionPolicy": "acknowledged_half_point_quantization",
        "ooxmlHalfPoints": 19,
        "effectiveOoxmlPoint": 9.5,
        "quantizationDeltaPoint": 0.2,
        "lineHeight": 1.52,
        "lineDxa": 365,
        "alignment": "justified",
        "characterSpacingPt": -0.2,
        "characterSpacingTwips": -4,
    }
    assert manifest["tables"] == [
        {
            "tableId": "TBL-01",
            "pageId": "P-01",
            "widthDxa": 8400,
            "columnWidthsDxa": [2400, 6000],
            "native": True,
        }
    ]
    assert manifest["figures"] == []
    assert manifest["artifacts"]["docx"]["path"] == str(
        result.generation / "document.docx"
    )
    assert manifest["artifacts"]["docx"]["sha256"] == _sha256(
        Path(manifest["artifacts"]["docx"]["path"])
    )
    assert result.docx == Path(request.output.docx_path)
    assert result.manifest == Path(request.output.manifest_path)
    assert result.docx.is_symlink()
    assert result.manifest.is_symlink()
    assert result.publication.is_symlink()
    assert result.publication.resolve(strict=True) == result.generation
    assert manifest["publication"] == {
        "pointerPath": str(result.publication),
        "generationPath": str(result.generation),
        "docxMember": "document.docx",
        "manifestMember": "manifest.json",
        "readerContract": "resolve_pointer_once_then_open_both_members",
        "compatibilityPathsAuthoritative": False,
    }
    assert manifest["inputs"]["pagePlanSha256"] == _canonical_sha(
        request.page_plan.model_dump(by_alias=True)
    )
    profile_sha = _canonical_sha(request.surface_profile.model_dump(by_alias=True))
    assert manifest["profile"] == {
        "profileId": "synthetic-korean-public-research-v1",
        "status": "locked",
        "sha256": profile_sha,
    }
    assert manifest["inputs"]["surfaceProfileSha256"] == profile_sha


def test_rejects_template_hash_drift_before_writing_outputs(tmp_path: Path) -> None:
    request = sample_request(tmp_path)
    request.template.sha256 = "0" * 64

    try:
        build_document(request)
    except ValueError as error:
        assert "template SHA-256" in str(error)
    else:
        raise AssertionError("template hash drift must block the build")

    assert not Path(request.output.docx_path).exists()
    assert not Path(request.output.manifest_path).exists()


def test_rejects_unrepresentable_exact_body_point_with_stable_rule_code(
    tmp_path: Path,
) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    payload["surfaceProfile"]["typography"]["precisionPolicy"] = "exact"

    with pytest.raises(ValidationError, match="KPP_TYPOGRAPHY_PRECISION"):
        BuildRequest.model_validate(payload)


def test_missing_precision_policy_defaults_to_exact_and_blocks_9_3pt(
    tmp_path: Path,
) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    del payload["surfaceProfile"]["typography"]["precisionPolicy"]

    with pytest.raises(ValidationError, match="KPP_TYPOGRAPHY_PRECISION"):
        BuildRequest.model_validate(payload)


def test_rejects_large_explicit_continuation_heading_without_override(
    tmp_path: Path,
) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    payload["pageArchitecture"] = _page_architecture(
        continuation=True,
        title_point_size=20.5,
    )
    with pytest.raises(
        ValidationError, match="KPP_PAGE_TITLE_CONTINUATION_LARGE"
    ):
        BuildRequest.model_validate(payload)


def test_allows_large_explicit_continuation_heading_with_bound_override(
    tmp_path: Path,
) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    payload["pageArchitecture"] = _page_architecture(
        continuation=True,
        title_point_size=20.5,
        issuer_override={
                "documentMode": "private_partnership",
                "modePolicyVersion": "1.0.0",
                "sourceId": "SRC-ISSUER-01",
                "reason": "파트너 지정 양식의 연속 면 제목 규칙",
        }
    )
    payload["issuerOverrideAuthorityIds"] = ["source:SRC-ISSUER-01"]

    request = BuildRequest.model_validate(payload)
    result = build_document(request)
    document_xml = _xml(result.docx, "word/document.xml")

    assert 'w:sz w:val="41"' in document_xml


def test_rejects_identity_bound_override_without_permitted_issuer_authority(
    tmp_path: Path,
) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    payload["pageArchitecture"] = _page_architecture(
        continuation=True,
        title_point_size=20.5,
        issuer_override={
            "documentMode": "private_partnership",
            "modePolicyVersion": "1.0.0",
            "sourceId": "SRC-ISSUER-01",
            "reason": "권한 검증이 필요한 예외",
        },
    )

    with pytest.raises(ValidationError, match="KPP_PAGE_ISSUER_OVERRIDE_UNBOUND"):
        BuildRequest.model_validate(payload)


def _page_architecture(
    *,
    continuation: bool,
    title_point_size: float,
    issuer_override: dict[str, object] | None = None,
) -> dict[str, object]:
    page: dict[str, object] = {
        "pageId": "P-01",
        "chapterId": "CH-01",
        "sectionId": "SEC-01",
        "pageRole": "research_method",
        "surfaceTemplateId": "r08-research-method-v1",
        "titleScope": "section",
        "titlePointSize": title_point_size,
        "continuation": continuation,
        "dominantSurface": "table",
        "surfaceVisibility": "internal",
        "claimIds": ["CLM-01"],
        "proofIds": [],
        "referenceIds": ["EV-01"],
        "figureIds": [],
    }
    if issuer_override is not None:
        page["issuerOverride"] = issuer_override
    return {
        "schemaVersion": "2.0.0",
        "projectId": "synthetic-research-proposal",
        "documentMode": "private_partnership",
        "modePolicyVersion": "1.0.0",
        "architectureStatus": "complete",
        "chapters": [{"chapterId": "CH-01"}],
        "sections": [{"sectionId": "SEC-01", "chapterId": "CH-01"}],
        "pages": [page],
    }


def test_manifest_publication_failure_rolls_back_docx_and_manifest(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    request = sample_request(tmp_path)
    docx_path = Path(request.output.docx_path)
    manifest_path = Path(request.output.manifest_path)
    real_replace = os.replace

    def fail_manifest_publish(
        source: str | os.PathLike[str],
        target: str | os.PathLike[str],
    ) -> None:
        if Path(target) == manifest_path:
            raise OSError("forced manifest publication failure")
        real_replace(source, target)

    monkeypatch.setattr(build_module.os, "replace", fail_manifest_publish)

    with pytest.raises(OSError, match="forced manifest publication failure"):
        build_document(request)

    assert not docx_path.exists()
    assert not manifest_path.exists()
    assert list(tmp_path.glob(".*.tmp")) == []
    assert list(tmp_path.glob(".*.backup")) == []


def test_publication_boundary_never_exposes_a_partial_artifact_pair(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    request = sample_request(tmp_path)
    docx_path = Path(request.output.docx_path)
    manifest_path = Path(request.output.manifest_path)
    real_replace = os.replace
    observed_publication: Path | None = None

    class SimulatedInterruption(RuntimeError):
        pass

    def interrupt_on_partial_pair(
        source: str | os.PathLike[str],
        target: str | os.PathLike[str],
    ) -> None:
        nonlocal observed_publication
        real_replace(source, target)
        publications = list(tmp_path.glob(".kpp-build-*/current"))
        for publication in publications:
            if not publication.exists():
                continue
            generation = publication.resolve(strict=True)
            published_docx = generation / "document.docx"
            published_manifest = generation / "manifest.json"
            assert published_docx.is_file()
            assert published_manifest.is_file()
            manifest = json.loads(published_manifest.read_text(encoding="utf-8"))
            assert manifest["artifacts"]["docx"]["sha256"] == _sha256(
                published_docx
            )
            observed_publication = publication
            raise SimulatedInterruption("interrupted immediately after publication")

        visible = (docx_path.exists(), manifest_path.exists())
        if visible[0] != visible[1]:
            raise SimulatedInterruption(f"observer saw partial artifact pair: {visible}")

    monkeypatch.setattr(build_module.os, "replace", interrupt_on_partial_pair)

    with pytest.raises(SimulatedInterruption):
        build_document(request)

    assert observed_publication is not None
    generation = observed_publication.resolve(strict=True)
    assert (generation / "document.docx").is_file()
    assert (generation / "manifest.json").is_file()


def test_republication_switches_one_complete_canonical_generation(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    request = sample_request(tmp_path)
    first = build_document(request)
    old_generation = first.publication.resolve(strict=True)
    payload = request.model_dump(by_alias=True)
    payload["projectId"] = "synthetic-research-proposal-revision"
    revised_request = BuildRequest.model_validate(payload)
    real_replace = os.replace
    observed_generations: list[Path] = []

    class SimulatedInterruption(RuntimeError):
        pass

    def interrupt_after_pointer_switch(
        source: str | os.PathLike[str],
        target: str | os.PathLike[str],
    ) -> None:
        real_replace(source, target)
        generation = first.publication.resolve(strict=True)
        published_docx = generation / "document.docx"
        published_manifest = generation / "manifest.json"
        manifest = json.loads(published_manifest.read_text(encoding="utf-8"))
        assert manifest["artifacts"]["docx"]["sha256"] == _sha256(published_docx)
        observed_generations.append(generation)
        if Path(target) == first.publication and generation != old_generation:
            raise SimulatedInterruption("interrupted after canonical pointer switch")

    monkeypatch.setattr(build_module.os, "replace", interrupt_after_pointer_switch)

    with pytest.raises(SimulatedInterruption):
        build_document(revised_request)

    new_generation = first.publication.resolve(strict=True)
    assert new_generation != old_generation
    assert set(observed_generations).issubset({old_generation, new_generation})
    assert (new_generation / "document.docx").is_file()
    assert (new_generation / "manifest.json").is_file()


def test_old_generation_manifest_keeps_its_document_after_republication(
    tmp_path: Path,
) -> None:
    request = sample_request(tmp_path)
    first = build_document(request)
    old_generation = first.generation
    old_manifest_path = old_generation / "manifest.json"
    old_manifest = json.loads(old_manifest_path.read_text(encoding="utf-8"))

    payload = request.model_dump(by_alias=True)
    payload["contentBlocks"][0]["paragraphs"][0]["text"] = (
        "개정된 자료와 현장 검증을 연결하여 후속 문서를 발행한다."
    )
    second = build_document(BuildRequest.model_validate(payload))

    old_manifest_after_republish = json.loads(
        old_manifest_path.read_text(encoding="utf-8")
    )
    old_document_path = Path(
        old_manifest_after_republish["artifacts"]["docx"]["path"]
    )
    assert second.generation != old_generation
    assert old_manifest_after_republish == old_manifest
    assert old_document_path == old_generation / "document.docx"
    assert _sha256(old_document_path) == old_manifest["artifacts"]["docx"]["sha256"]
    assert "개정된 자료" not in "\n".join(
        paragraph.text for paragraph in Document(old_document_path).paragraphs
    )


def test_rejects_preexisting_generations_symlink_without_publishing(
    tmp_path: Path,
) -> None:
    request = sample_request(tmp_path)
    docx_path = Path(request.output.docx_path)
    manifest_path = Path(request.output.manifest_path)
    publication = build_module._publication_path(docx_path, manifest_path)
    bundle_root = publication.parent
    bundle_root.mkdir()
    alternate = bundle_root / "alternate"
    alternate.mkdir()
    (bundle_root / "generations").symlink_to(alternate.name, target_is_directory=True)

    with pytest.raises(
        ValueError,
        match="publication generations directory must not be a symlink",
    ):
        build_document(request)

    assert not docx_path.exists()
    assert not manifest_path.exists()
    assert not publication.exists()
    assert list(alternate.iterdir()) == []


def test_manifest_directory_probe_leaves_no_partial_docx(tmp_path: Path) -> None:
    request = sample_request(tmp_path)
    docx_path = Path(request.output.docx_path)
    manifest_path = Path(request.output.manifest_path)
    manifest_path.mkdir()

    with pytest.raises(IsADirectoryError, match="output target is not a file"):
        build_document(request)

    assert not docx_path.exists()
    assert manifest_path.is_dir()
    assert list(tmp_path.glob(".*.tmp")) == []


def test_failed_republication_preserves_previous_artifact_pair(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    request = sample_request(tmp_path)
    docx_path = Path(request.output.docx_path)
    manifest_path = Path(request.output.manifest_path)
    docx_path.write_bytes(b"previous docx")
    manifest_path.write_text("previous manifest\n", encoding="utf-8")
    real_replace = os.replace
    manifest_publish_failed = False

    def fail_manifest_publish(
        source: str | os.PathLike[str],
        target: str | os.PathLike[str],
    ) -> None:
        nonlocal manifest_publish_failed
        if Path(target) == manifest_path and not manifest_publish_failed:
            manifest_publish_failed = True
            raise OSError("forced manifest publication failure")
        real_replace(source, target)

    monkeypatch.setattr(build_module.os, "replace", fail_manifest_publish)

    with pytest.raises(OSError, match="forced manifest publication failure"):
        build_document(request)

    assert docx_path.read_bytes() == b"previous docx"
    assert manifest_path.read_text(encoding="utf-8") == "previous manifest\n"


def test_build_embeds_only_hash_bound_figure_on_its_planned_page(tmp_path: Path) -> None:
    # A deterministic 1x1 transparent PNG; the builder applies the locked Word width.
    figure = tmp_path / "figure.png"
    figure.write_bytes(
        b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
            "AQUBAScY42YAAAAASUVORK5CYII="
        )
    )
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    payload["figureManifest"]["figures"] = [
        {
            "figureId": "FIG-01",
            "requirementId": "REQ-01",
            "pageId": "P-01",
            "claimIds": ["CLM-01"],
            "renderer": "svg-flow",
            "path": str(figure),
            "sha256": _sha256(figure),
            "format": "png",
            "caption": "그림 1. 연구 수행 구조",
            "evidenceIds": ["EV-01"],
            "widthDxa": 3600,
        }
    ]
    payload["contentBlocks"][0]["figureIds"] = ["FIG-01"]
    payload["pagePlan"]["pages"][0]["figureSpecs"] = [
        {
            "figureId": "FIG-01",
            "requirementId": "REQ-01",
            "pageId": "P-01",
            "title": "연구 수행 구조",
            "intent": "flow",
            "dataShape": "process_flow",
            "decisionTask": "연구 수행 구조를 확인한다.",
            "semanticValueIntent": "causal_mechanism",
            "decisionEffect": "연구 수행 단계의 인과관계를 판단한다.",
            "nonDuplicateOf": ["BLK-RESEARCH-NARRATIVE"],
            "encodedVariables": ["input", "process", "output"],
            "claimIds": ["CLM-01"],
            "evidenceIds": ["EV-01"],
            "family": "flow",
            "renderer": "svg-flow",
        }
    ]

    result = build_document(BuildRequest.model_validate(payload))
    document_xml = _xml(result.docx, "word/document.xml")
    manifest = json.loads(result.manifest.read_text(encoding="utf-8"))

    assert "<w:drawing>" in document_xml
    assert 'wp:extent cx="2286000"' in document_xml
    assert manifest["figures"] == [
        {
            "caption": "그림 1. 연구 수행 구조",
            "claimIds": ["CLM-01"],
            "embedded": True,
            "evidenceIds": ["EV-01"],
            "figureId": "FIG-01",
            "format": "png",
            "pageId": "P-01",
            "path": str(figure.resolve()),
            "renderer": "svg-flow",
            "requirementId": "REQ-01",
            "sha256": _sha256(figure),
        }
    ]


def test_build_request_accepts_decorative_figure_with_zero_evidentiary_bindings(
    tmp_path: Path,
) -> None:
    figure = tmp_path / "decorative.png"
    figure.write_bytes(
        b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
            "AQUBAScY42YAAAAASUVORK5CYII="
        )
    )
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    payload["figureManifest"]["figures"] = [{
        "figureId": "FIG-DECORATIVE-01",
        "requirementId": "REQ-01",
        "pageId": "P-01",
        "claimIds": [],
        "renderer": "svg-flow",
        "path": str(figure),
        "sha256": _sha256(figure),
        "format": "png",
        "caption": "그림 1. 장식용 구분 표지",
        "evidenceIds": [],
        "widthDxa": 3600,
    }]
    payload["contentBlocks"][0]["figureIds"] = ["FIG-DECORATIVE-01"]
    payload["pagePlan"]["pages"][0]["figureSpecs"] = [{
        "figureId": "FIG-DECORATIVE-01",
        "requirementId": "REQ-01",
        "pageId": "P-01",
        "title": "장식용 구분 표지",
        "intent": "flow",
        "dataShape": "process_flow",
        "decisionTask": "장 구분을 보조한다.",
        "semanticValueIntent": "decorative",
        "decisionEffect": "",
        "nonDuplicateOf": [],
        "encodedVariables": [],
        "claimIds": [],
        "evidenceIds": [],
        "family": "flow",
        "renderer": "svg-flow",
    }]

    result = build_document(BuildRequest.model_validate(payload))

    assert result.docx.exists()


def test_build_request_rejects_unknown_schema_version(tmp_path: Path) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    payload["schemaVersion"] = "2.0.0"

    with pytest.raises(ValidationError):
        BuildRequest.model_validate(payload)


@pytest.mark.parametrize(
    ("mutate", "message"),
    [
        (
            lambda payload: payload["evidenceLedger"]["claims"][0].update(
                evidenceIds=["EV-UNKNOWN"]
            ),
            "claim evidenceIds must exist",
        ),
        (
            lambda payload: payload["evidenceLedger"]["bindings"][0].update(
                targetPageRole="wrong_role"
            ),
            "page role must match",
        ),
        (
            lambda payload: payload["evidenceLedger"]["claims"][0].update(
                evidenceIds=[]
            ),
            "links must be reciprocal",
        ),
    ],
)
def test_rejects_broken_evidence_page_cross_references(
    tmp_path: Path,
    mutate: object,
    message: str,
) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    mutate(payload)

    with pytest.raises(ValidationError, match=message):
        BuildRequest.model_validate(payload)


def test_rejects_cross_claim_paragraph_evidence_pairing(tmp_path: Path) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    payload["pagePlan"]["pages"][0]["claimIds"].append("CLM-02")
    payload["evidenceLedger"]["claims"].append(
        {"claimId": "CLM-02", "status": "verified", "evidenceIds": ["EV-02"]}
    )
    payload["evidenceLedger"]["bindings"].append(
        {
            "evidenceId": "EV-02",
            "sourcePath": "/locked/other-source.pdf",
            "sourceSha256": "b" * 64,
            "scope": "다른 주장에 대한 공식 근거",
            "claimIds": ["CLM-02"],
            "targetRequirementId": "REQ-01",
            "targetPageId": "P-01",
            "targetPageRole": "research_method",
        }
    )
    payload["contentBlocks"][0]["paragraphs"][0].update(
        claimIds=["CLM-01"],
        evidenceIds=["EV-02"],
    )

    with pytest.raises(
        ValidationError,
        match="paragraph evidenceIds must support every paragraph claimId",
    ):
        BuildRequest.model_validate(payload)


def test_rejects_cross_claim_figure_evidence_pairing(tmp_path: Path) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    payload["pagePlan"]["pages"][0]["claimIds"].append("CLM-02")
    payload["evidenceLedger"]["claims"].append(
        {"claimId": "CLM-02", "status": "verified", "evidenceIds": ["EV-02"]}
    )
    payload["evidenceLedger"]["bindings"].append(
        {
            "evidenceId": "EV-02",
            "sourcePath": "/locked/other-source.pdf",
            "sourceSha256": "b" * 64,
            "scope": "다른 주장에 대한 공식 근거",
            "claimIds": ["CLM-02"],
            "targetRequirementId": "REQ-01",
            "targetPageId": "P-01",
            "targetPageRole": "research_method",
        }
    )
    payload["figureManifest"]["figures"] = [
        {
            "figureId": "FIG-01",
            "requirementId": "REQ-01",
            "pageId": "P-01",
            "claimIds": ["CLM-01"],
            "renderer": "svg-flow",
            "path": str(tmp_path / "not-read-during-validation.png"),
            "sha256": "c" * 64,
            "format": "png",
            "caption": "그림 1. 연구 수행 구조",
            "evidenceIds": ["EV-02"],
            "widthDxa": 3600,
        }
    ]
    payload["contentBlocks"][0]["figureIds"] = ["FIG-01"]
    payload["pagePlan"]["pages"][0]["figureSpecs"] = [
        {
            "figureId": "FIG-01",
            "requirementId": "REQ-01",
            "pageId": "P-01",
            "title": "연구 수행 구조",
            "intent": "flow",
            "dataShape": "process_flow",
            "decisionTask": "연구 수행 구조를 확인한다.",
            "semanticValueIntent": "causal_mechanism",
            "decisionEffect": "연구 수행 단계의 인과관계를 판단한다.",
            "nonDuplicateOf": ["BLK-RESEARCH-NARRATIVE"],
            "encodedVariables": ["input", "process", "output"],
            "claimIds": ["CLM-01"],
            "evidenceIds": ["EV-02"],
            "family": "flow",
            "renderer": "svg-flow",
        }
    ]

    with pytest.raises(
        ValidationError,
        match="figure evidenceIds must support every figure claimId",
    ):
        BuildRequest.model_validate(payload)


def test_rejects_figure_evidence_bound_to_another_page(tmp_path: Path) -> None:
    payload = sample_request(tmp_path).model_dump(by_alias=True)
    second_page = {
        **payload["pagePlan"]["pages"][0],
        "pageId": "P-02",
        "requirementId": "REQ-02",
        "pageRole": "expected_effect",
    }
    payload["pagePlan"]["pages"].append(second_page)
    payload["contentBlocks"].append(
        {
            **payload["contentBlocks"][0],
            "pageId": "P-02",
            "paragraphs": [],
            "tables": [],
            "figureIds": ["FIG-01"],
        }
    )
    payload["figureManifest"]["figures"] = [
        {
            "figureId": "FIG-01",
            "requirementId": "REQ-02",
            "pageId": "P-02",
            "claimIds": ["CLM-01"],
            "renderer": "svg-flow",
            "path": str(tmp_path / "not-read-during-validation.png"),
            "sha256": "b" * 64,
            "format": "png",
            "caption": "그림 1. 기대효과",
            "evidenceIds": ["EV-01"],
            "widthDxa": 3600,
        }
    ]
    payload["pagePlan"]["pages"][1]["figureSpecs"] = [
        {
            "figureId": "FIG-01",
            "requirementId": "REQ-02",
            "pageId": "P-02",
            "title": "기대효과",
            "intent": "flow",
            "dataShape": "process_flow",
            "decisionTask": "기대효과를 확인한다.",
            "semanticValueIntent": "decision_tradeoff",
            "decisionEffect": "기대효과의 선택 기준과 결과를 판단한다.",
            "nonDuplicateOf": ["BLK-EFFECT-NARRATIVE"],
            "encodedVariables": ["option", "criterion", "effect"],
            "claimIds": ["CLM-01"],
            "evidenceIds": ["EV-01"],
            "family": "flow",
            "renderer": "svg-flow",
        }
    ]

    with pytest.raises(ValidationError, match="figure evidenceIds must target"):
        BuildRequest.model_validate(payload)
