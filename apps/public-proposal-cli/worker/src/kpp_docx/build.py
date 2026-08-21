"""Governed Word-native proposal document builder."""

from __future__ import annotations

import hashlib
import json
import os
import shutil
import tempfile
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Twips
from pydantic import BaseModel, ConfigDict, Field, model_validator

from .styles import (
    TypographyContract,
    format_body_paragraph,
    format_navigation_paragraph,
    format_run,
    install_governed_styles,
)
from .tables import TableContract, add_native_table


SCHEMA_VERSION = "1.0.0"
BUILDER_VERSION = "0.1.0"
SHA256_PATTERN = r"^[a-f0-9]{64}$"


class StrictModel(BaseModel):
    model_config = ConfigDict(extra="forbid", populate_by_name=True)


class KppBuildError(ValueError):
    """Build failure carrying a stable machine-readable rule code."""

    def __init__(self, code: str, message: str) -> None:
        super().__init__(f"{code}: {message}")
        self.code = code


class TemplateRef(StrictModel):
    asset_id: str = Field(alias="assetId", min_length=1)
    path: str = Field(min_length=1)
    sha256: str = Field(pattern=SHA256_PATTERN)


class PlannedFigureSpec(StrictModel):
    figure_id: str = Field(alias="figureId", min_length=1)
    requirement_id: str = Field(alias="requirementId", min_length=1)
    page_id: str = Field(alias="pageId", min_length=1)
    title: str = Field(min_length=1)
    intent: Literal[
        "schedule",
        "responsibility",
        "matrix",
        "comparison",
        "evidence_chain",
        "research_framework",
        "flow",
    ]
    data_shape: Literal[
        "time_axis",
        "responsibility_matrix",
        "two_by_two",
        "comparison_series",
        "evidence_links",
        "research_framework",
        "process_flow",
    ] = Field(alias="dataShape")
    decision_task: str = Field(alias="decisionTask", min_length=1)
    semantic_value_intent: Literal[
        "data_evidence",
        "causal_mechanism",
        "decision_tradeoff",
        "operational_control",
        "decorative",
    ] = Field(alias="semanticValueIntent")
    decision_effect: str = Field(alias="decisionEffect")
    non_duplicate_of: list[str] = Field(alias="nonDuplicateOf")
    encoded_variables: list[str] = Field(alias="encodedVariables")
    claim_ids: list[str] = Field(alias="claimIds")
    evidence_ids: list[str] = Field(alias="evidenceIds")
    family: Literal[
        "gantt",
        "raci",
        "matrix",
        "comparison_chart",
        "evidence_chain",
        "framework",
        "flow",
    ]
    renderer: Literal[
        "svg-gantt",
        "word-native-raci-table",
        "svg-raci-matrix",
        "svg-2x2-matrix",
        "svg-comparison-chart",
        "svg-evidence-chain",
        "svg-academic-framework",
        "svg-flow",
    ]

    @model_validator(mode="after")
    def validate_semantic_value_bindings(self) -> "PlannedFigureSpec":
        if self.semantic_value_intent == "decorative":
            if any((self.decision_effect, self.non_duplicate_of, self.encoded_variables, self.claim_ids, self.evidence_ids)):
                raise ValueError("decorative figures must not carry evidentiary bindings")
            return self
        if (not self.decision_effect.strip() or not self.non_duplicate_of or not self.encoded_variables
                or not self.claim_ids or not self.evidence_ids):
            raise ValueError("non-decorative figures require semantic value and evidentiary bindings")
        return self


class IssuerOverride(StrictModel):
    document_mode: Literal[
        "public_procurement",
        "research_service",
        "private_partnership",
        "internal_decision",
        "document_restyle",
    ] = Field(alias="documentMode")
    mode_policy_version: str = Field(alias="modePolicyVersion", min_length=1)
    rule_id: str | None = Field(alias="ruleId", min_length=1, default=None)
    source_id: str | None = Field(alias="sourceId", min_length=1, default=None)
    reason: str = Field(min_length=1)

    @model_validator(mode="after")
    def validate_locator(self) -> "IssuerOverride":
        if (self.rule_id is None) == (self.source_id is None):
            raise ValueError("issuerOverride must identify exactly one ruleId or sourceId")
        return self


class SurfaceRepetitionException(StrictModel):
    rule_id: Literal["issuer_mandatory_form", "accessibility_repeated_instruction"] = Field(alias="ruleId")
    source_id: str = Field(alias="sourceId", min_length=1)
    source_sha256: str = Field(alias="sourceSha256", pattern=SHA256_PATTERN)
    rationale: str = Field(min_length=1)


class FigureSpec(StrictModel):
    figure_id: str = Field(alias="figureId", min_length=1)
    requirement_id: str = Field(alias="requirementId", min_length=1)
    page_id: str = Field(alias="pageId", min_length=1)
    claim_ids: list[str] = Field(alias="claimIds")
    renderer: str = Field(min_length=1)
    path: str = Field(min_length=1)
    sha256: str = Field(pattern=SHA256_PATTERN)
    format: Literal["png", "jpeg", "jpg"]
    caption: str = Field(min_length=1)
    evidence_ids: list[str] = Field(alias="evidenceIds")
    width_dxa: int = Field(alias="widthDxa", gt=0, default=7200)


class FigureManifest(StrictModel):
    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    figures: list[FigureSpec]


class PagePlanItem(StrictModel):
    page_id: str = Field(alias="pageId", min_length=1)
    requirement_id: str = Field(alias="requirementId", min_length=1)
    page_role: str = Field(alias="pageRole", min_length=1)
    surface_template_id: str = Field(alias="surfaceTemplateId", min_length=1)
    claim_ids: list[str] = Field(alias="claimIds")
    figure_specs: list[PlannedFigureSpec] = Field(alias="figureSpecs")


class PageArchitectureItem(StrictModel):
    page_id: str = Field(alias="pageId", min_length=1)
    chapter_id: str = Field(alias="chapterId", min_length=1)
    section_id: str = Field(alias="sectionId", min_length=1)
    page_role: str = Field(alias="pageRole", min_length=1)
    surface_template_id: str = Field(alias="surfaceTemplateId", min_length=1)
    title_scope: Literal["cover", "chapter", "section", "surface", "none"] = Field(
        alias="titleScope"
    )
    title_point_size: float | None = Field(alias="titlePointSize", gt=0, le=72, default=None)
    continuation: bool
    dominant_surface: Literal["narrative", "table", "figure", "mixed", "form"] = Field(
        alias="dominantSurface"
    )
    surface_visibility: Literal["internal", "reader"] = Field(alias="surfaceVisibility")
    evaluation_question: str | None = Field(alias="evaluationQuestion", min_length=1, default=None)
    direct_answer: str | None = Field(alias="directAnswer", min_length=1, default=None)
    claim_ids: list[str] = Field(alias="claimIds")
    proof_ids: list[str] = Field(alias="proofIds")
    reference_ids: list[str] = Field(alias="referenceIds")
    figure_ids: list[str] = Field(alias="figureIds")
    continuity_from_page_id: str | None = Field(alias="continuityFromPageId", min_length=1, default=None)
    continuity_to_page_id: str | None = Field(alias="continuityToPageId", min_length=1, default=None)
    issuer_override: IssuerOverride | None = Field(alias="issuerOverride", default=None)
    surface_repetition_exception: SurfaceRepetitionException | None = Field(alias="surfaceRepetitionException", default=None)


class PageArchitecture(StrictModel):
    schema_version: str = Field(alias="schemaVersion", min_length=1)
    project_id: str = Field(alias="projectId", min_length=1)
    document_mode: Literal[
        "public_procurement",
        "research_service",
        "private_partnership",
        "internal_decision",
        "document_restyle",
    ] = Field(alias="documentMode")
    mode_policy_version: str = Field(alias="modePolicyVersion", min_length=1)
    architecture_status: Literal["staged", "complete"] = Field(alias="architectureStatus")
    chapters: list[dict[str, object]]
    sections: list[dict[str, object]]
    pages: list[PageArchitectureItem] = Field(min_length=1)


class PagePlan(StrictModel):
    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    pages: list[PagePlanItem] = Field(min_length=1)


class EvidenceClaim(StrictModel):
    claim_id: str = Field(alias="claimId", min_length=1)
    status: Literal["verified", "bounded", "pending_blank", "blocked"]
    evidence_ids: list[str] = Field(alias="evidenceIds")


class EvidenceBinding(StrictModel):
    evidence_id: str = Field(alias="evidenceId", min_length=1)
    source_path: str = Field(alias="sourcePath", min_length=1)
    source_sha256: str = Field(alias="sourceSha256", pattern=SHA256_PATTERN)
    scope: str = Field(min_length=1)
    claim_ids: list[str] = Field(alias="claimIds", min_length=1)
    target_requirement_id: str = Field(alias="targetRequirementId", min_length=1)
    target_page_id: str = Field(alias="targetPageId", min_length=1)
    target_page_role: str = Field(alias="targetPageRole", min_length=1)


class EvidenceLedger(StrictModel):
    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    claims: list[EvidenceClaim]
    bindings: list[EvidenceBinding]


class ParagraphContent(StrictModel):
    text: str
    claim_ids: list[str] = Field(alias="claimIds")
    evidence_ids: list[str] = Field(alias="evidenceIds")


class TableContent(StrictModel):
    table_id: str = Field(alias="tableId", min_length=1)
    caption: str = Field(min_length=1)
    headers: list[str] = Field(min_length=1)
    rows: list[list[str]]
    column_widths_dxa: list[int] = Field(alias="columnWidthsDxa", min_length=1)


class ContentBlock(StrictModel):
    page_id: str = Field(alias="pageId", min_length=1)
    heading: str = Field(min_length=1)
    paragraphs: list[ParagraphContent]
    tables: list[TableContent]
    figure_ids: list[str] = Field(alias="figureIds")


class TypographyProfile(StrictModel):
    heading_font: Literal["Noto Sans CJK KR"] = Field(alias="headingFont")
    navigation_font: Literal["Noto Sans CJK KR"] = Field(alias="navigationFont")
    body_font: Literal["Noto Serif CJK KR"] = Field(alias="bodyFont")
    body_point: Literal[9.3] = Field(alias="bodyPoint")
    line_height: Literal[1.52] = Field(alias="lineHeight")
    alignment: Literal["justified"]
    character_spacing_pt: float = Field(alias="characterSpacingPt", ge=-2, le=2)
    precision_policy: Literal["exact", "acknowledged_half_point_quantization"] = Field(
        alias="precisionPolicy",
        default="exact",
    )

    @model_validator(mode="after")
    def validate_ooxml_precision(self) -> "TypographyProfile":
        effective_point = round(self.body_point * 2) / 2
        if effective_point != self.body_point and self.precision_policy == "exact":
            raise KppBuildError(
                "KPP_TYPOGRAPHY_PRECISION",
                f"{self.body_point}pt cannot be represented in OOXML half-points",
            )
        return self


class CellMargins(StrictModel):
    top: int = Field(ge=0)
    start: int = Field(ge=0)
    bottom: int = Field(ge=0)
    end: int = Field(ge=0)


class TableProfile(StrictModel):
    width_dxa: int = Field(alias="widthDxa", gt=0)
    cell_margin_dxa: CellMargins = Field(alias="cellMarginDxa")
    border_size_eighth_pt: int = Field(alias="borderSizeEighthPt", gt=0)


class SurfaceProfile(StrictModel):
    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    profile_id: str = Field(alias="profileId", min_length=1)
    status: Literal["locked"]
    typography: TypographyProfile
    table: TableProfile


class BuildOutput(StrictModel):
    docx_path: str = Field(alias="docxPath", min_length=1)
    manifest_path: str = Field(alias="manifestPath", min_length=1)


class BuildRequest(StrictModel):
    """Versioned, closed input boundary for one document build."""

    schema_version: Literal[SCHEMA_VERSION] = Field(alias="schemaVersion")
    project_id: str = Field(alias="projectId", min_length=1)
    template: TemplateRef
    page_plan: PagePlan = Field(alias="pagePlan")
    page_architecture: PageArchitecture | None = Field(alias="pageArchitecture", default=None)
    issuer_override_authority_ids: list[str] = Field(
        alias="issuerOverrideAuthorityIds", default_factory=list
    )
    evidence_ledger: EvidenceLedger = Field(alias="evidenceLedger")
    content_blocks: list[ContentBlock] = Field(alias="contentBlocks", min_length=1)
    figure_manifest: FigureManifest = Field(alias="figureManifest")
    surface_profile: SurfaceProfile = Field(alias="surfaceProfile")
    output: BuildOutput

    @model_validator(mode="after")
    def validate_cross_references(self) -> "BuildRequest":
        page_ids = [page.page_id for page in self.page_plan.pages]
        if len(page_ids) != len(set(page_ids)):
            raise ValueError("pagePlan pageId values must be unique")
        content_page_ids = [block.page_id for block in self.content_blocks]
        if len(content_page_ids) != len(set(content_page_ids)):
            raise ValueError("content block pageId values must be unique")
        if set(content_page_ids) != set(page_ids):
            raise ValueError("content blocks must map exactly to pagePlan pages")

        if self.page_architecture is not None:
            architecture = self.page_architecture
            architecture_page_ids = [page.page_id for page in architecture.pages]
            if architecture.project_id != self.project_id:
                raise KppBuildError(
                    "KPP_PAGE_ARCHITECTURE_IDENTITY",
                    "pageArchitecture projectId must match BuildRequest",
                )
            if architecture_page_ids != page_ids:
                raise KppBuildError(
                    "KPP_PAGE_ARCHITECTURE_PAGES",
                    "pageArchitecture pages must match pagePlan order exactly",
                )
            for page in architecture.pages:
                override = page.issuer_override
                override_bound = (
                    override is not None
                    and override.document_mode == architecture.document_mode
                    and override.mode_policy_version == architecture.mode_policy_version
                    and _issuer_override_authority_id(override)
                    in self.issuer_override_authority_ids
                )
                if override is not None and not override_bound:
                    raise KppBuildError(
                        "KPP_PAGE_ISSUER_OVERRIDE_UNBOUND",
                        "issuerOverride must match pageArchitecture mode identity",
                    )
                point_size = _architecture_title_point_size(page)
                if page.continuation and point_size > 12 and not override_bound:
                    raise KppBuildError(
                        "KPP_PAGE_TITLE_CONTINUATION_LARGE",
                        f"{page.page_id} continuation heading is "
                        f"{point_size}pt; maximum is 12pt",
                    )

        figures = {figure.figure_id: figure for figure in self.figure_manifest.figures}
        if len(figures) != len(self.figure_manifest.figures):
            raise ValueError("figure manifest figureId values must be unique")
        referenced_figures = {
            figure_id for block in self.content_blocks for figure_id in block.figure_ids
        }
        if referenced_figures != set(figures):
            raise ValueError("content figureIds must map exactly to the figure manifest")
        if any(figure.page_id not in page_ids for figure in figures.values()):
            raise ValueError("figure pageId must reference a planned page")

        planned_figures: dict[str, PlannedFigureSpec] = {}
        for page in self.page_plan.pages:
            for figure in page.figure_specs:
                if figure.figure_id in planned_figures:
                    raise ValueError("pagePlan figureId values must be unique")
                if figure.page_id != page.page_id:
                    raise ValueError("planned figure pageId must match its planned page")
                if figure.requirement_id != page.requirement_id:
                    raise ValueError("planned figure requirementId must match its planned page")
                if not set(figure.claim_ids).issubset(set(page.claim_ids)):
                    raise ValueError("planned figure claimIds must be declared by its planned page")
                planned_figures[figure.figure_id] = figure
        for figure_id, figure in figures.items():
            planned = planned_figures.get(figure_id)
            if planned is None:
                raise ValueError("figure manifest entries must reference a planned figureSpec")
            if (
                figure.requirement_id != planned.requirement_id
                or figure.page_id != planned.page_id
                or figure.claim_ids != planned.claim_ids
                or figure.evidence_ids != planned.evidence_ids
                or figure.renderer != planned.renderer
            ):
                raise ValueError("figure manifest entries must match their planned figureSpec")

        evidence_ids = [binding.evidence_id for binding in self.evidence_ledger.bindings]
        if len(evidence_ids) != len(set(evidence_ids)):
            raise ValueError("evidence ledger evidenceId values must be unique")
        claim_ids = [claim.claim_id for claim in self.evidence_ledger.claims]
        if len(claim_ids) != len(set(claim_ids)):
            raise ValueError("evidence ledger claimId values must be unique")
        known_evidence_ids = set(evidence_ids)
        known_claim_ids = set(claim_ids)

        claim_by_id = {claim.claim_id: claim for claim in self.evidence_ledger.claims}
        binding_by_id = {
            binding.evidence_id: binding for binding in self.evidence_ledger.bindings
        }
        for claim in self.evidence_ledger.claims:
            if not set(claim.evidence_ids).issubset(known_evidence_ids):
                raise ValueError("claim evidenceIds must exist in the evidence ledger")
            if any(
                claim.claim_id not in binding_by_id[evidence_id].claim_ids
                for evidence_id in claim.evidence_ids
            ):
                raise ValueError("claim and evidence binding links must be reciprocal")
        for binding in self.evidence_ledger.bindings:
            if not set(binding.claim_ids).issubset(known_claim_ids):
                raise ValueError("evidence binding claimIds must exist in the evidence ledger")
            if any(
                binding.evidence_id not in claim_by_id[claim_id].evidence_ids
                for claim_id in binding.claim_ids
            ):
                raise ValueError("claim and evidence binding links must be reciprocal")

        page_by_id = {page.page_id: page for page in self.page_plan.pages}
        for page in self.page_plan.pages:
            if not set(page.claim_ids).issubset(known_claim_ids):
                raise ValueError("page claimIds must exist in the evidence ledger")
        for binding in self.evidence_ledger.bindings:
            target_page = page_by_id.get(binding.target_page_id)
            if target_page is None:
                raise ValueError("evidence binding targetPageId must reference a planned page")
            if binding.target_requirement_id != target_page.requirement_id:
                raise ValueError("evidence binding requirement must match its planned page")
            if binding.target_page_role != target_page.page_role:
                raise ValueError("evidence binding page role must match its planned page")
            if not set(binding.claim_ids).issubset(set(target_page.claim_ids)):
                raise ValueError("evidence binding claimIds must be declared by its target page")

        for block in self.content_blocks:
            planned_claim_ids = set(page_by_id[block.page_id].claim_ids)
            for paragraph in block.paragraphs:
                if not set(paragraph.claim_ids).issubset(planned_claim_ids):
                    raise ValueError("paragraph claimIds must be declared by its planned page")
                if not set(paragraph.evidence_ids).issubset(known_evidence_ids):
                    raise ValueError("paragraph evidenceIds must exist in the evidence ledger")
                if any(
                    binding_by_id[evidence_id].target_page_id != block.page_id
                    for evidence_id in paragraph.evidence_ids
                ):
                    raise ValueError("paragraph evidenceIds must target its planned page")
                if any(
                    not set(paragraph.claim_ids).issubset(
                        set(binding_by_id[evidence_id].claim_ids)
                    )
                    for evidence_id in paragraph.evidence_ids
                ):
                    raise ValueError(
                        "paragraph evidenceIds must support every paragraph claimId"
                    )
            for figure_id in block.figure_ids:
                figure = figures[figure_id]
                if figure.page_id != block.page_id:
                    raise ValueError("figure pageId must match the content block pageId")
                if not set(figure.evidence_ids).issubset(known_evidence_ids):
                    raise ValueError("figure evidenceIds must exist in the evidence ledger")
                if any(
                    binding_by_id[evidence_id].target_page_id != block.page_id
                    for evidence_id in figure.evidence_ids
                ):
                    raise ValueError("figure evidenceIds must target its planned page")
                if any(
                    not set(figure.claim_ids).issubset(
                        set(binding_by_id[evidence_id].claim_ids)
                    )
                    for evidence_id in figure.evidence_ids
                ):
                    raise ValueError(
                        "figure evidenceIds must support every figure claimId"
                    )

        table_width = self.surface_profile.table.width_dxa
        for block in self.content_blocks:
            for table in block.tables:
                if len(table.column_widths_dxa) != len(table.headers):
                    raise ValueError("table widths must match the header count")
                if sum(table.column_widths_dxa) != table_width:
                    raise ValueError("table widths must sum to the locked profile width")
                if any(len(row) != len(table.headers) for row in table.rows):
                    raise ValueError("table rows must match the header count")
        return self


class BuildResult(BaseModel):
    """Build paths plus the canonical snapshot reader entrypoint.

    Atomic readers must resolve ``publication`` exactly once, then open
    ``document.docx`` and ``manifest.json`` inside that immutable generation.
    ``docx`` and ``manifest`` retain the requested compatibility paths but are
    not a snapshot boundary when a later build may publish concurrently.
    """

    model_config = ConfigDict(arbitrary_types_allowed=True)
    docx: Path
    manifest: Path
    publication: Path
    generation: Path


def build_document(request: BuildRequest) -> BuildResult:
    """Build a proposal DOCX and a manifest bound to its exact inputs and output."""

    template_path = Path(request.template.path).expanduser().resolve()
    if not template_path.is_file():
        raise ValueError(f"template is missing: {template_path}")
    template_hash = _sha256_file(template_path)
    if template_hash != request.template.sha256:
        raise ValueError("template SHA-256 does not match the locked BuildRequest")

    figure_records: list[dict[str, object]] = []
    for figure in request.figure_manifest.figures:
        figure_path = Path(figure.path).expanduser().resolve()
        if not figure_path.is_file():
            raise ValueError(f"figure is missing: {figure_path}")
        if _sha256_file(figure_path) != figure.sha256:
            raise ValueError(f"figure SHA-256 does not match: {figure.figure_id}")
        figure_records.append(
            {
                "figureId": figure.figure_id,
                "requirementId": figure.requirement_id,
                "pageId": figure.page_id,
                "claimIds": figure.claim_ids,
                "renderer": figure.renderer,
                "path": str(figure_path),
                "sha256": figure.sha256,
                "format": figure.format,
                "caption": figure.caption,
                "evidenceIds": figure.evidence_ids,
                "embedded": True,
            }
        )

    docx_path = _output_path(request.output.docx_path)
    manifest_path = _output_path(request.output.manifest_path)
    if docx_path == manifest_path:
        raise ValueError("DOCX and manifest output paths must be distinct")
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document(template_path)
    _clear_document_body(document)

    typography = _typography_contract(request.surface_profile.typography)
    style_ids = install_governed_styles(document, typography)
    furniture = _document_furniture(request)
    for section in document.sections:
        if section.header.paragraphs and section.header.paragraphs[0].runs:
            section.header.paragraphs[0].runs[0].text = furniture["header"]
        if section.footer.paragraphs and section.footer.paragraphs[0].runs:
            # Replace only the label run so the PAGE field remains intact.
            section.footer.paragraphs[0].runs[0].text = furniture["footer"]
        for paragraph in section.header.paragraphs:
            format_navigation_paragraph(paragraph, typography)
        for paragraph in section.footer.paragraphs:
            format_navigation_paragraph(paragraph, typography)
    table_contract = _table_contract(request.surface_profile.table)
    content_by_page_id = {block.page_id: block for block in request.content_blocks}
    architecture_by_page_id = {
        page.page_id: page for page in request.page_architecture.pages
    } if request.page_architecture is not None else {}
    figure_by_id = {figure.figure_id: figure for figure in request.figure_manifest.figures}
    table_records: list[dict[str, object]] = []

    for block_index, planned_page in enumerate(request.page_plan.pages):
        block = content_by_page_id[planned_page.page_id]
        if block_index:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        architecture_page = architecture_by_page_id.get(planned_page.page_id)
        # A chapter opener may introduce the title.  Continuation pages can
        # explicitly choose titleScope=none so the narrative continues without
        # a repeated page-title shell; the compact section heading remains
        # available for issuers that explicitly request titleScope=section.
        if architecture_page is None or architecture_page.title_scope != "none":
            heading = document.add_paragraph(style=style_ids["heading"])
            heading_run = heading.add_run(block.heading)
            title_point_size = (
                _architecture_title_point_size(architecture_page)
                if architecture_page is not None
                else 16
            )
            format_run(
                heading_run,
                font=typography.heading_font,
                half_points=round(title_point_size * 2),
                bold=True,
            )
        for paragraph_content in block.paragraphs:
            paragraph = document.add_paragraph()
            paragraph.add_run(paragraph_content.text)
            format_body_paragraph(paragraph, typography)

        for table_content in block.tables:
            caption = document.add_paragraph(style=style_ids["caption"])
            caption_run = caption.add_run(table_content.caption)
            format_run(
                caption_run,
                font=typography.heading_font,
                half_points=18,
                bold=True,
            )
            add_native_table(
                document,
                headers=table_content.headers,
                rows=table_content.rows,
                column_widths_dxa=table_content.column_widths_dxa,
                contract=table_contract,
                typography=typography,
            )
            table_records.append(
                {
                    "tableId": table_content.table_id,
                    "pageId": block.page_id,
                    "widthDxa": table_contract.width_dxa,
                    "columnWidthsDxa": table_content.column_widths_dxa,
                    "native": True,
                }
            )

        for figure_id in block.figure_ids:
            figure = figure_by_id[figure_id]
            caption = document.add_paragraph(style=style_ids["caption"])
            caption_run = caption.add_run(figure.caption)
            format_run(
                caption_run,
                font=typography.heading_font,
                half_points=18,
                bold=True,
            )
            document.add_picture(
                str(Path(figure.path).expanduser().resolve()),
                width=Twips(figure.width_dxa),
            )

    staged_docx: Path | None = None
    staged_manifest: Path | None = None
    try:
        staged_docx = _sibling_temp_path(docx_path)
        staged_manifest = _sibling_temp_path(manifest_path)
        document.save(staged_docx)
        _fsync_file(staged_docx)
        docx_hash = _sha256_file(staged_docx)
        publication_path = _publication_path(docx_path, manifest_path)
        generation_path = _generation_path(
            publication_path,
            docx_hash=docx_hash,
            request=request,
        )

        effective_point = typography.body_half_points / 2
        manifest = {
            "schemaVersion": SCHEMA_VERSION,
            "builderVersion": BUILDER_VERSION,
            "projectId": request.project_id,
            "template": {
                "assetId": request.template.asset_id,
                "path": str(template_path),
                "sha256": template_hash,
            },
            "profile": {
                "profileId": request.surface_profile.profile_id,
                "status": request.surface_profile.status,
                "sha256": _sha256_json(request.surface_profile),
            },
            "inputs": {
                "pagePlanSha256": _sha256_json(request.page_plan),
                "evidenceLedgerSha256": _sha256_json(request.evidence_ledger),
                "contentBlocksSha256": _sha256_json(request.content_blocks),
                "figureManifestSha256": _sha256_json(request.figure_manifest),
                "surfaceProfileSha256": _sha256_json(request.surface_profile),
            },
            "pages": [
                {
                    "pageId": page.page_id,
                    "pageRole": page.page_role,
                    "surfaceTemplateId": page.surface_template_id,
                }
                for page in request.page_plan.pages
            ],
            "styles": {
                "heading": {"font": typography.heading_font},
                "navigation": {"font": typography.navigation_font},
                "body": {
                    "font": typography.body_font,
                    "requestedPoint": typography.body_point,
                    "precisionPolicy": request.surface_profile.typography.precision_policy,
                    "ooxmlHalfPoints": typography.body_half_points,
                    "effectiveOoxmlPoint": effective_point,
                    "quantizationDeltaPoint": round(
                        effective_point - typography.body_point, 3
                    ),
                    "lineHeight": typography.line_height,
                    "lineDxa": typography.body_line_dxa,
                    "alignment": "justified",
                    "characterSpacingPt": typography.character_spacing_pt,
                    "characterSpacingTwips": typography.character_spacing_twips,
                },
            },
            "tables": table_records,
            "figures": figure_records,
            "artifacts": {
                "docx": {
                    "path": str(generation_path / "document.docx"),
                    "sha256": docx_hash,
                },
            },
            "publication": {
                "pointerPath": str(publication_path),
                "generationPath": str(generation_path),
                "docxMember": "document.docx",
                "manifestMember": "manifest.json",
                "readerContract": "resolve_pointer_once_then_open_both_members",
                "compatibilityPathsAuthoritative": False,
            },
        }
        with staged_manifest.open("w", encoding="utf-8") as stream:
            stream.write(
                json.dumps(manifest, ensure_ascii=False, sort_keys=True, indent=2)
                + "\n"
            )
            stream.flush()
            os.fsync(stream.fileno())
        _publish_artifact_pair(
            ((staged_docx, docx_path), (staged_manifest, manifest_path)),
            publication=publication_path,
            generation=generation_path,
        )
    finally:
        if staged_docx is not None:
            staged_docx.unlink(missing_ok=True)
        if staged_manifest is not None:
            staged_manifest.unlink(missing_ok=True)
    return BuildResult(
        docx=docx_path,
        manifest=manifest_path,
        publication=publication_path,
        generation=generation_path,
    )


def _clear_document_body(document: object) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _typography_contract(profile: TypographyProfile) -> TypographyContract:
    return TypographyContract(
        heading_font=profile.heading_font,
        navigation_font=profile.navigation_font,
        body_font=profile.body_font,
        body_point=profile.body_point,
        line_height=profile.line_height,
        character_spacing_pt=profile.character_spacing_pt,
    )


def _architecture_title_point_size(page: PageArchitectureItem) -> float:
    if page.title_point_size is not None:
        return page.title_point_size
    return 20.5 if page.title_scope in ("cover", "chapter") else 12


def _document_furniture(request: BuildRequest) -> dict[str, str]:
    mode = request.page_architecture.document_mode if request.page_architecture is not None else None
    if mode is None:
        first_role = request.page_plan.pages[0].page_role
        if first_role in {"research_method", "research_question", "evidence_plan", "limitations", "utilization_plan"}:
            mode = "research_service"
        elif first_role in {"mutual_value", "party_roles", "operating_model", "collaboration_options", "next_decision"}:
            mode = "private_partnership"
        elif first_role in {"decision_request", "alternatives", "tradeoffs", "risk_register", "owner_approval"}:
            mode = "internal_decision"
        elif first_role in {"source_inventory", "content_ledger", "mutation_report", "layout_accessibility", "acceptance_record"}:
            mode = "document_restyle"
        else:
            mode = "public_procurement"
    mode_label = {
        "public_procurement": "공공조달 제안서",
        "research_service": "연구용역 제안서",
        "private_partnership": "민간협력 제안서",
        "internal_decision": "내부 의사결정 문서",
        "document_restyle": "문서 재구성 검증본",
    }[mode]
    return {
        "header": f"{mode_label} | 형식·구조 검증본",
        "footer": f"{request.content_blocks[0].heading} | ",
    }


def _issuer_override_authority_id(override: IssuerOverride) -> str:
    if override.source_id is not None:
        return f"source:{override.source_id}"
    return f"rule:{override.rule_id}"


def _table_contract(profile: TableProfile) -> TableContract:
    return TableContract(
        width_dxa=profile.width_dxa,
        margin_top_dxa=profile.cell_margin_dxa.top,
        margin_start_dxa=profile.cell_margin_dxa.start,
        margin_bottom_dxa=profile.cell_margin_dxa.bottom,
        margin_end_dxa=profile.cell_margin_dxa.end,
        border_size_eighth_pt=profile.border_size_eighth_pt,
    )


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _sha256_json(value: object) -> str:
    if isinstance(value, BaseModel):
        value = value.model_dump(by_alias=True)
    elif isinstance(value, list):
        value = [
            item.model_dump(by_alias=True) if isinstance(item, BaseModel) else item
            for item in value
        ]
    encoded = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _output_path(raw_path: str) -> Path:
    """Canonicalize an output parent without following its final alias."""

    path = Path(raw_path).expanduser()
    if not path.is_absolute():
        path = Path.cwd() / path
    return path.parent.resolve() / path.name


def _sibling_temp_path(target: Path) -> Path:
    descriptor, raw_path = tempfile.mkstemp(
        prefix=f".{target.name}.",
        suffix=".tmp",
        dir=target.parent,
    )
    os.close(descriptor)
    return Path(raw_path)


def _fsync_file(path: Path) -> None:
    with path.open("rb") as stream:
        os.fsync(stream.fileno())


def _fsync_directories(paths: set[Path]) -> None:
    for path in paths:
        descriptor = os.open(path, os.O_RDONLY)
        try:
            os.fsync(descriptor)
        finally:
            os.close(descriptor)


def _publication_path(docx_path: Path, manifest_path: Path) -> Path:
    pair_key = hashlib.sha256(
        f"{docx_path}\0{manifest_path}".encode("utf-8")
    ).hexdigest()[:16]
    return docx_path.parent / f".kpp-build-{pair_key}" / "current"


def _generation_path(
    publication: Path,
    *,
    docx_hash: str,
    request: BuildRequest,
) -> Path:
    request_hash = _sha256_json(request)
    generation_id = hashlib.sha256(
        f"{docx_hash}\0{request_hash}".encode("ascii")
    ).hexdigest()
    return publication.parent / "generations" / generation_id


def _publish_artifact_pair(
    pairs: tuple[tuple[Path, Path], ...],
    *,
    publication: Path,
    generation: Path,
) -> None:
    """Publish an immutable pair behind one atomically switched pointer.

    ``publication`` is the only snapshot boundary. Compatibility output paths
    are stable aliases for existing callers, but concurrent readers must
    resolve the publication pointer once and read both generation members.
    """

    if len(pairs) != 2:
        raise ValueError("artifact publication requires exactly two members")
    bundle_root = publication.parent
    generations = bundle_root / "generations"
    bundle_parent = bundle_root.parent.resolve()
    if bundle_root.is_symlink():
        raise ValueError(
            f"publication bundle root must not be a symlink: {bundle_root}"
        )
    bundle_root.mkdir(parents=True, exist_ok=True)
    if generations.is_symlink():
        raise ValueError(
            "publication generations directory must not be a symlink: "
            f"{generations}"
        )
    generations.mkdir(exist_ok=True)
    if bundle_root.resolve().parent != bundle_parent:
        raise ValueError("publication bundle escaped its output parent")
    if generations.resolve().parent != bundle_root.resolve():
        raise ValueError("publication generations directory escaped its bundle")
    if generation.parent != generations or generation.name in {"", ".", ".."}:
        raise ValueError(
            "generation path must be contained by the publication bundle"
        )

    backups: list[tuple[Path, Path]] = []
    installed_aliases: list[Path] = []
    temporary_aliases: list[Path] = []
    parent_paths = {target.parent for _, target in pairs}
    temporary_generation = Path(
        tempfile.mkdtemp(prefix=".generation-", dir=generations)
    )
    temporary_pointer: Path | None = None
    generation_created = False
    pointer_replace_started = False
    try:
        for _, target in pairs:
            if target.exists() and not target.is_file():
                raise IsADirectoryError(f"output target is not a file: {target}")

        member_names = ("document.docx", "manifest.json")
        for (staged, _), member_name in zip(pairs, member_names, strict=True):
            os.replace(staged, temporary_generation / member_name)
        _fsync_directories({temporary_generation})

        if generation.exists():
            if not generation.is_dir():
                raise ValueError(f"generation target is not a directory: {generation}")
            for member_name in member_names:
                existing = generation / member_name
                candidate = temporary_generation / member_name
                if not existing.is_file() or _sha256_file(existing) != _sha256_file(
                    candidate
                ):
                    raise ValueError("existing immutable generation content differs")
            shutil.rmtree(temporary_generation)
        else:
            os.replace(temporary_generation, generation)
            generation_created = True
            _fsync_directories({generations})

        for (_, target), member_name in zip(pairs, member_names, strict=True):
            expected_target = publication / member_name
            if target.is_symlink() and Path(os.readlink(target)) == expected_target:
                continue
            if target.exists() or target.is_symlink():
                descriptor, raw_backup = tempfile.mkstemp(
                    prefix=f".{target.name}.",
                    suffix=".backup",
                    dir=target.parent,
                )
                os.close(descriptor)
                backup = Path(raw_backup)
                backup.unlink()
                os.replace(target, backup)
                backups.append((target, backup))

            descriptor, raw_alias = tempfile.mkstemp(
                prefix=f".{target.name}.",
                suffix=".tmp",
                dir=target.parent,
            )
            os.close(descriptor)
            alias = Path(raw_alias)
            temporary_aliases.append(alias)
            alias.unlink()
            os.symlink(expected_target, alias)
            os.replace(alias, target)
            installed_aliases.append(target)
        _fsync_directories(parent_paths)

        descriptor, raw_pointer = tempfile.mkstemp(
            prefix=".current-",
            suffix=".tmp",
            dir=bundle_root,
        )
        os.close(descriptor)
        temporary_pointer = Path(raw_pointer)
        temporary_pointer.unlink()
        os.symlink(Path("generations") / generation.name, temporary_pointer)
        _fsync_directories({bundle_root})
        pointer_replace_started = True
        os.replace(temporary_pointer, publication)
        if not _publication_resolves_to(publication, generation):
            raise ValueError("published generation pointer escaped its bundle")
        _fsync_directories({bundle_root})
    except Exception:
        committed = pointer_replace_started and _publication_resolves_to(
            publication, generation
        )
        if not committed:
            for target in reversed(installed_aliases):
                target.unlink(missing_ok=True)
            for target, backup in reversed(backups):
                if backup.exists():
                    os.replace(backup, target)
            if generation_created:
                shutil.rmtree(generation, ignore_errors=True)
            _fsync_directories(parent_paths | {generations})
        raise
    finally:
        if temporary_pointer is not None:
            temporary_pointer.unlink(missing_ok=True)
        for alias in temporary_aliases:
            alias.unlink(missing_ok=True)
        shutil.rmtree(temporary_generation, ignore_errors=True)
        for _, backup in backups:
            backup.unlink(missing_ok=True)


def _publication_resolves_to(publication: Path, generation: Path) -> bool:
    if not publication.is_symlink():
        return False
    try:
        resolved = publication.resolve(strict=True)
    except (FileNotFoundError, RuntimeError):
        return False
    generations = publication.parent / "generations"
    return (
        resolved == generation.resolve()
        and resolved.parent == generations.resolve()
    )
