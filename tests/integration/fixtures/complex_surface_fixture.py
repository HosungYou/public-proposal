"""Build a small but multi-surface research-service document fixture.

The fixture is deliberately generated from structured values so the surface
contract can test real OOXML/SVG/render bytes without depending on a project
checkout or a private proposal artifact.
"""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Final

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image


CONTRACT: Final[dict[str, object]] = {
    "schemaVersion": "kpp-surface-contract-1.0",
    "requireRenderManifest": True,
    "tables": {
        "headerFill": "#E8EEF5",
        "bodyFill": "#FFFFFF",
        "repeatHeader": True,
        "headerAlignment": "center",
        "bodyAlignment": "left",
        "bodyLine": {"line": "365", "lineRule": "auto"},
        "allowZebraStriping": False,
    },
    "svg": {
        "allowOuterCanvasFill": False,
        "bodyFill": "#FFFFFF",
        "rowRoles": ["work-package-row", "raci-row"],
        "allowZebraStriping": False,
    },
    "render": {"requirePages": 4},
}


def materialize(root: Path, variant: str = "valid") -> dict[str, Path]:
    root.mkdir(parents=True, exist_ok=True)
    docx_path = root / "research-service.docx"
    svg_dir = root / "figures"
    pages_dir = root / "pages"
    svg_dir.mkdir(exist_ok=True)
    pages_dir.mkdir(exist_ok=True)

    _write_document(docx_path)
    svg_paths = [
        _write_framework(svg_dir / "framework.svg"),
        _write_gantt(svg_dir / "gantt.svg"),
        _write_raci(svg_dir / "raci.svg"),
        _write_flow(svg_dir / "flow.svg"),
    ]
    page_paths = []
    for page_no in range(1, 5):
        page_path = pages_dir / f"page-{page_no}.png"
        Image.new("RGB", (210, 297), "white").save(page_path)
        page_paths.append(page_path)
    page_hashes = {path: sha256(path) for path in page_paths}

    if variant == "missing_table_shading":
        _remove_first_header_shading(docx_path)
    elif variant == "outer_canvas_fill":
        svg_paths[0].write_text(
            svg_paths[0].read_text(encoding="utf-8").replace(
                "><title>",
                '><rect width="720" height="376" fill="#FCFCFA"/><title>',
                1,
            ),
            encoding="utf-8",
        )
    elif variant == "zebra_rows":
        svg_paths[1].write_text(
            svg_paths[1].read_text(encoding="utf-8").replace('fill="#FFFFFF"', 'fill="#F4F6F8"', 1),
            encoding="utf-8",
        )
    elif variant == "stale_figure_hash":
        svg_paths[2].write_text(
            svg_paths[2].read_text(encoding="utf-8").replace("역할과 승인", "역할과 승인 변경", 1),
            encoding="utf-8",
        )
    elif variant == "missing_page":
        page_paths[-1].unlink()

    contract_path = root / "surface-contract.json"
    contract_path.write_text(json.dumps(CONTRACT, ensure_ascii=False, indent=2), encoding="utf-8")
    manifest_path = root / "render-manifest.json"
    manifest_path.write_text(
        json.dumps(
            {
                "schemaVersion": "kpp-render-observation-1.0",
                "docx": {"path": docx_path.name, "sha256": sha256(docx_path)},
                "figures": [{"path": str(path.relative_to(root)), "sha256": sha256(path)} for path in svg_paths],
                "pages": [
                    {"pageNo": index, "path": str(path.relative_to(root)), "sha256": page_hashes[path]}
                    for index, path in enumerate(page_paths, 1)
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    if variant == "stale_figure_hash":
        svg_paths[2].write_text(svg_paths[2].read_text(encoding="utf-8") + "<!-- stale bytes -->\n", encoding="utf-8")
    return {
        "docx": docx_path,
        "svg_dir": svg_dir,
        "manifest": manifest_path,
        "contract": contract_path,
    }


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _write_document(path: Path) -> None:
    document = Document()
    tables = [
        (
            ["평가기준", "직답", "근거", "상태"],
            [["연구 설계", "혼합방법", "EV-RS-01", "확인"], ["확산성", "현장 적용", "EV-RS-02", "보완"]],
        ),
        (
            ["Claim ID", "Proof ID", "산출물", "검수 기준", "담당"],
            [["CLM-RS-01", "EV-RS-01", "중간보고서", "교차검증", "연구책임자"], ["CLM-RS-02", "EV-RS-02", "현장 매뉴얼", "사용성 확인", "운영책임자"]],
        ),
        (
            ["과업", "착수", "분석", "검토", "승인"],
            [["W1 문헌 검토", "R", "A", "C", "I"], ["W2 현장 검증", "C", "R", "A", "I"]],
        ),
    ]
    for index, (headers, rows) in enumerate(tables):
        document.add_heading(
            ["평가기준 대응 교차표", "근거·산출물 추적 대장", "연구 수행·승인 RACI"][index],
            level=1,
        )
        _add_table(document, headers, rows)
        document.add_paragraph("표의 판단 경계와 후속 조치는 다음 검토 단계에서 확인한다.")
    document.save(path)


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header)
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell(cell, value, header=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            _set_cell(cell, value, header=False)


def _set_cell(cell, value: str, *, header: bool) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "E8EEF5" if header else "FFFFFF")
    cell._tc.get_or_add_tcPr().append(shading)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    ppr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "365")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    paragraph.add_run(value)


def _remove_first_header_shading(path: Path) -> None:
    document = Document(path)
    cell = document.tables[0].rows[0].cells[0]
    shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    assert shading is not None
    shading.getparent().remove(shading)
    document.save(path)


def _svg(title: str, body: str) -> str:
    return f'<svg xmlns="http://www.w3.org/2000/svg" width="720" height="376"><title>{title}</title><desc>근거 기반 복합 표면</desc><style>.title{{font-size:12pt}}.meta{{font-size:8pt}}</style><text class="title" x="24" y="30">{title}</text>{body}</svg>\n'


def _write_framework(path: Path) -> Path:
    path.write_text(_svg("연구 프레임워크", '<rect x="32" y="80" width="180" height="90" fill="#F4F6F8" stroke="#7D8894"/><rect x="260" y="80" width="180" height="90" fill="#F4F6F8" stroke="#7D8894"/><path d="M212 125 H260" stroke="#082F63"/>'), encoding="utf-8")
    return path


def _write_gantt(path: Path) -> Path:
    path.write_text(_svg("연구 일정", '<g data-kpp-role="work-package-row"><rect x="24" y="80" width="672" height="56" fill="#FFFFFF"/><text x="32" y="110">W1 문헌 검토</text></g><g data-kpp-role="work-package-row"><rect x="24" y="136" width="672" height="56" fill="#FFFFFF"/><text x="32" y="166">W2 현장 검증</text></g>'), encoding="utf-8")
    return path


def _write_raci(path: Path) -> Path:
    path.write_text(_svg("역할과 승인", '<g data-kpp-role="raci-row"><rect x="24" y="80" width="672" height="56" fill="#FFFFFF"/><text x="32" y="110">R1 연구책임</text></g><g data-kpp-role="raci-row"><rect x="24" y="136" width="672" height="56" fill="#FFFFFF"/><text x="32" y="166">R2 운영책임</text></g>'), encoding="utf-8")
    return path


def _write_flow(path: Path) -> Path:
    path.write_text(_svg("검토 의사결정 흐름", '<rect x="80" y="100" width="200" height="70" fill="#F4F6F8" stroke="#7D8894"/><rect x="400" y="100" width="200" height="70" fill="#E8EEF5" stroke="#7D8894"/><path d="M280 135 H400" stroke="#082F63"/>'), encoding="utf-8")
    return path


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("root", type=Path)
    parser.add_argument("--variant", default="valid")
    arguments = parser.parse_args()
    print(json.dumps({key: str(value) for key, value in materialize(arguments.root, arguments.variant).items()}, ensure_ascii=False))
